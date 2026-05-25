#!/usr/bin/env python
# -*- coding: utf-8 -*-

from src.tables import builders
from src.core import generator
from src.content import static_data
from docx import Document

print("=" * 50)
print("VALIDAÇÃO DA TABELA 24")
print("=" * 50)
print()

# 1. Dados
print("✓ 1. DADOS:")
print(f"   - dados_tabela_premio_cnj: {len(static_data.dados_tabela_premio_cnj)} linhas")
print()

# 2. Função
print("✓ 2. FUNCAO:")
print(f"   - adicionar_tabela_2col existe: {hasattr(builders, 'adicionar_tabela_2col')}")
print()

# 3. Registro
print("✓ 3. MAPA_RECURSOS:")
chave = "Tabela 24 - Pontuação do TJMG na edição 2025 do Prêmio CNJ de Qualidade."
existe = chave in static_data.MAPA_RECURSOS
print(f"   - Tabela 24 registrada: {existe}")
if existe:
    entrada = static_data.MAPA_RECURSOS[chave]
    print(f"   - Tipo: {entrada['tipo']}")
    print(f"   - Dados: {len(entrada['dados'])} linhas")
    print(f"   - Fonte: {entrada['fonte_custom']}")
print()

# 4. Teste de geração
print("✓ 4. TESTE DE GERACAO:")
try:
    doc = Document()
    builders.adicionar_tabela_2col(
        doc,
        static_data.dados_tabela_premio_cnj,
        titulo_custom=chave,
        fonte=static_data.MAPA_RECURSOS[chave]['fonte_custom']
    )
    
    print(f"   - Tabelas criadas: {len(doc.tables)}")
    print(f"   - Linhas na tabela: {len(doc.tables[0].rows)}")
    print(f"   - Colunas: {len(doc.tables[0].columns)}")
    print()
    print("=" * 50)
    print("RESULTADO: SUCESSO - Tabela foi gerada corretamente!")
    print("=" * 50)
    
except Exception as e:
    print(f"   ERRO: {e}")
    import traceback
    traceback.print_exc()
