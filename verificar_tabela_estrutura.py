#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def verificar_tabelas():
    """Verifica a estrutura das tabelas no documento gerado"""
    doc = Document('teste_metas_institucionais.docx')
    
    print("=" * 80)
    print("VERIFICAÇÃO DE ESTRUTURA DAS TABELAS")
    print("=" * 80)
    
    tabelas = doc.tables
    print(f"\n✓ Total de tabelas encontradas: {len(tabelas)}\n")
    
    if len(tabelas) > 0:
        # Verifica primeira tabela
        tabela = tabelas[0]
        print(f"Analisando primeira tabela:")
        print(f"  Linhas: {len(tabela.rows)}")
        print(f"  Colunas: {len(tabela.columns)}")
        
        print(f"\nConteúdo da primeira tabela:")
        for row_idx, row in enumerate(tabela.rows):
            conteudo = []
            for cell_idx, cell in enumerate(row.cells):
                conteudo.append(cell.text.strip()[:30])  # primeiros 30 caracteres
            print(f"  Linha {row_idx + 1}: {' | '.join(conteudo)}")
        
        # Verifica cores
        print(f"\nCores da primeira tabela (primeira linha):")
        for cell_idx, cell in enumerate(tabela.rows[0].cells):
            tcPr = cell._element.tcPr
            if tcPr is not None:
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    print(f"  Coluna {cell_idx + 1}: {fill}")
        
        print(f"\nCores da segunda tabela (segunda linha):")
        for cell_idx, cell in enumerate(tabela.rows[1].cells):
            tcPr = cell._element.tcPr
            if tcPr is not None:
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    print(f"  Coluna {cell_idx + 1}: {fill}")
        
        # Verifica título
        print(f"\nVerificando títulos e formatação:")
        paragrafos = doc.paragraphs[:20]  # primeiros 20 parágrafos
        
        for p_idx, p in enumerate(paragrafos):
            if p.text.startswith('TJMG'):
                print(f"\n  Parágrafo {p_idx}: {p.text}")
                print(f"    Estilo: {p.style.name}")
                if len(p.runs) > 0:
                    run = p.runs[0]
                    print(f"    Cor: {run.font.color.rgb if run.font.color.rgb else 'padrão'}")
                    print(f"    Tamanho: {run.font.size}")
    
    print("\n" + "=" * 80)
    print("✓ VERIFICAÇÃO CONCLUÍDA")
    print("=" * 80)

if __name__ == '__main__':
    verificar_tabelas()
