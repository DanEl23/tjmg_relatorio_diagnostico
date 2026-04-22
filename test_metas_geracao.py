#!/usr/bin/env python3
"""
Teste completo de geração de METAS_INSTITUCIONAIS no documento
"""

import sys
sys.path.insert(0, '.')

from docx import Document
from docx.shared import Pt

print("=" * 70)
print("TESTE COMPLETO: GERAÇÃO DE METAS_INSTITUCIONAIS")
print("=" * 70)

try:
    # 1. Cria um documento de teste
    print("\n1. Criando documento de teste...")
    doc_teste = Document()
    print("   ✓ Documento criado")
    
    # 2. Importa a função processar_recurso
    print("\n2. Importando função processar_recurso...")
    from src.core.generator import processar_recurso
    from src.content.static_data import MAPA_RECURSOS
    print("   ✓ Função e MAPA_RECURSOS importados")
    
    # 3. Obtém a entrada de METAS_INSTITUCIONAIS do MAPA_RECURSOS
    print("\n3. Obtendo configuração de METAS_INSTITUCIONAIS...")
    config_metas = MAPA_RECURSOS.get("METAS_INSTITUCIONAIS")
    if not config_metas:
        raise ValueError("METAS_INSTITUCIONAIS não encontrado em MAPA_RECURSOS")
    print(f"   ✓ Configuração obtida: tipo={config_metas.get('tipo')}")
    
    # 4. Chama processar_recurso
    print("\n4. Chamando processar_recurso com METAS_INSTITUCIONAIS...")
    processar_recurso(doc_teste, "METAS_INSTITUCIONAIS", config_metas, loader_jn=None)
    print("   ✓ processar_recurso executado com sucesso")
    
    # 5. Verifica o documento gerado
    print("\n5. Analisando documento gerado...")
    total_paragrafos = len(doc_teste.paragraphs)
    total_tabelas = len(doc_teste.tables)
    
    print(f"   ✓ Parágrafos adicionados: {total_paragrafos}")
    print(f"   ✓ Tabelas adicionadas: {total_tabelas}")
    
    # 6. Mostra amostra do conteúdo
    print("\n6. Amostra do conteúdo gerado:")
    for i, para in enumerate(doc_teste.paragraphs[:15]):
        if para.text.strip():
            texto_amostra = para.text[:70]
            if len(para.text) > 70:
                texto_amostra += "..."
            print(f"   [{i}] {texto_amostra}")
    
    # 7. Salva em arquivo de teste
    print("\n7. Salvando documento de teste...")
    arquivo_saida = "teste_metas_institucionais.docx"
    doc_teste.save(arquivo_saida)
    print(f"   ✓ Documento salvo em: {arquivo_saida}")
    
    print("\n" + "=" * 70)
    print("✓ TESTE COMPLETO PASSOU COM SUCESSO!")
    print("=" * 70)
    
except Exception as e:
    print(f"\n✗ ERRO: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
