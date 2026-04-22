#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def verificar_marcador_no_document():
    """Verifica se o marcador de metas está no Conteudo_Fonte"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("VERIFICAÇÃO: MARCADOR EM CONTEUDO_FONTE.docx")
    print("=" * 100)
    
    print(f"\nTotal de parágrafos: {len(doc.paragraphs)}")
    
    # Procura pelo marcador
    encontrou = False
    for idx, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if "META" in texto and "INSTITUCIONAL" in texto:
            encontrou = True
            print(f"\n[Parágrafo {idx}] {texto}")
    
    if not encontrou:
        print("\n❌ Marcador METAS_INSTITUCIONAIS não encontrado!")
    else:
        print("\n✓ Marcador encontrado!")
    
    # Mostra parágrafos próximos à conclusão
    print("\n" + "=" * 100)
    print("PARÁGRAFOS PRÓXIMOS À CONCLUSÃO:")
    print("=" * 100)
    
    for idx, p in enumerate(doc.paragraphs[-30:]):
        texto = p.text.strip()
        real_idx = len(doc.paragraphs) - 30 + idx
        if texto:
            print(f"[{real_idx}] {texto[:80]}")

if __name__ == '__main__':
    verificar_marcador_no_document()
