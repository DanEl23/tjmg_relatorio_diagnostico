#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para verificar se as formatações das tabelas estão corretas
no documento gerado de metas institucionais.
"""

from docx import Document
from docx.oxml.ns import qn

# Abre o documento gerado
doc = Document('teste_metas_institucionais.docx')

print("=" * 70)
print("VERIFICAÇÃO DE FORMATAÇÃO DAS TABELAS")
print("=" * 70)

# Lista todas as tabelas
num_tabelas = len(doc.tables)
print(f"\n✓ Total de tabelas encontradas: {num_tabelas}")

# Inspeciona a primeira tabela como exemplo
if num_tabelas > 0:
    print(f"\nAnalisando primeira tabela (exemplo):")
    print("-" * 70)
    
    table = doc.tables[0]
    print(f"Número de linhas: {len(table.rows)}")
    print(f"Número de colunas: {len(table.columns)}")
    
    # Cores esperadas
    cores_esperadas = {
        'Linha 1 (Header)': 'RGB(89,89,89) - Cinza escuro',
        'Linha 2 (Data)': 'RGB(166,166,166) - Cinza médio',
        'Linha 3 (Header Anos)': 'RGB(89,89,89) - Cinza escuro',
        'Linha 4 (Data Anos)': 'RGB(231,231,231) - Cinza claro',
    }
    
    for linha_idx, (desc, cor) in enumerate(cores_esperadas.items()):
        if linha_idx < len(table.rows):
            print(f"\n{desc}:")
            print(f"  Esperado: {cor}")
            
            row = table.rows[linha_idx]
            for col_idx, cell in enumerate(row.cells):
                # Extrai cor de fundo
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    shading = tcPr.find(qn('w:shd'))
                    if shading is not None:
                        fill = shading.get(qn('w:fill'))
                        altura = row.height
                        print(f"  Coluna {col_idx}: Fill={fill}, Altura={altura}")
                    else:
                        print(f"  Coluna {col_idx}: Sem cor de fundo detectada")
    
    print("\n" + "=" * 70)
    print("DIMENSÕES DAS COLUNAS E LINHAS")
    print("-" * 70)
    
    # Verifica largura das colunas (esperado: 3cm = ~1134 twips)
    print(f"Colunas esperadas: 6")
    print(f"Largura esperada por coluna: 3cm (~1701 twips)")
    print(f"Altura esperada por linha: 0.6cm")
    
    row = table.rows[0]
    for col_idx, cell in enumerate(row.cells):
        tcPr = cell._element.tcPr
        if tcPr is not None:
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                width_twips = tcW.get(qn('w:w'))
                width_cm = float(width_twips) / 1701 * 3 if width_twips else 0
                print(f"  Coluna {col_idx}: {width_twips} twips (~{width_cm:.2f}cm)")
            else:
                print(f"  Coluna {col_idx}: Sem largura específica")
        else:
            print(f"  Coluna {col_idx}: Sem tcPr")
    
    for row_idx, row in enumerate(table.rows[:4]):
        altura = row.height
        if altura:
            altura_cm = altura / 914285 * 2.54  # Conversão aproximada
            print(f"  Linha {row_idx}: {altura} EMU (~{altura_cm:.2f}cm)")
    
    print("\n" + "=" * 70)
    print("BORDAS DAS CÉLULAS")
    print("-" * 70)
    
    primeira_celula = table.rows[0].cells[0]
    tcPr = primeira_celula._element.tcPr
    if tcPr is not None:
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            for border_elem in tcBorders:
                border_type = border_elem.tag.split('}')[1]
                size = border_elem.get(qn('w:sz'))
                color = border_elem.get(qn('w:color'))
                print(f"  Borda {border_type}: Size={size} (1/2pt=4), Color={color}")
        else:
            print("  ⚠ Bordas não detectadas")

print("\n" + "=" * 70)
print("✓ VERIFICAÇÃO CONCLUÍDA")
print("=" * 70)
