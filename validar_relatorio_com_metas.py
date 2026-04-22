#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def validar_relatorio_com_metas():
    """Valida se o relatório contém as 55 metas institucionais"""
    
    doc = Document('data/output/Relatorio_Com_Metas_2026.docx')
    
    print("=" * 100)
    print("VALIDAÇÃO: RELATORIO_COM_METAS_2026.docx")
    print("=" * 100)
    
    print(f"\nTotal de parágrafos: {len(doc.paragraphs)}")
    print(f"Total de tabelas: {len(doc.tables)}")
    
    # Procura pelas metas institucionais
    metas_encontradas = []
    
    for idx, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if texto.startswith('TJMG '):
            # Extrai número da meta
            partes = texto.split('-')
            if len(partes) > 0:
                num_meta = partes[0].strip()
                metas_encontradas.append(num_meta)
    
    print(f"\n✓ METAS ENCONTRADAS: {len(metas_encontradas)}")
    
    if len(metas_encontradas) > 0:
        print(f"\nPrimeiras 10 metas:")
        for i, meta in enumerate(metas_encontradas[:10]):
            print(f"  {i+1}. {meta}")
        
        if len(metas_encontradas) > 10:
            print(f"\n  ...")
            print(f"\nÚltimas 5 metas:")
            for i, meta in enumerate(metas_encontradas[-5:], start=len(metas_encontradas)-4):
                print(f"  {i}. {meta}")
    
    print(f"\n{'='*100}")
    if len(metas_encontradas) >= 55:
        print(f"✓✓✓ SUCESSO! {len(metas_encontradas)} METAS INSTITUCIONAIS INSERIDAS NO RELATÓRIO!")
    else:
        print(f"⚠️  AVISO: Esperado 55 metas, encontrado {len(metas_encontradas)}")
    print(f"{'='*100}")
    
    # Procura pela seção "METAS INSTITUCIONAIS"
    secao_encontrada = False
    for p in doc.paragraphs:
        if "METAS INSTITUCIONAIS" in p.text.upper():
            secao_encontrada = True
            break
    
    if secao_encontrada:
        print("✓ Seção 'METAS INSTITUCIONAIS' encontrada")
    else:
        print("❌ Seção 'METAS INSTITUCIONAIS' NÃO encontrada")
    
    print(f"\nArquivo: data/output/Relatorio_Com_Metas_2026.docx")

if __name__ == '__main__':
    validar_relatorio_com_metas()
