#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import RGBColor

def verificar_formatacao_detalhada():
    """Verifica a formatação detalhada das tabelas no documento gerado"""
    doc = Document('teste_metas_institucionais.docx')
    
    print("=" * 100)
    print("VERIFICAÇÃO DETALHADA DE FORMATAÇÃO DAS TABELAS")
    print("=" * 100)
    
    tabelas = doc.tables
    print(f"\n✓ Total de tabelas encontradas: {len(tabelas)}\n")
    
    if len(tabelas) > 0:
        # Analisa primeira tabela
        tabela = tabelas[0]
        print(f"Analisando primeira tabela (TJMG 5):\n")
        
        # Linha 1 - Cabeçalho
        print("=" * 100)
        print("LINHA 1 - CABEÇALHO (Ano | 2022 | 2023 | 2024 | 2025)")
        print("=" * 100)
        
        for col_idx, cell in enumerate(tabela.rows[0].cells):
            print(f"\nColuna {col_idx + 1}:")
            print(f"  Texto: {cell.text}")
            
            # Cor de fundo
            tcPr = cell._element.tcPr
            if tcPr is not None:
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    print(f"  Cor de fundo: {fill}")
            
            # Verificar negrito, tamanho e cor da fonte
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    print(f"  Negrito: {run.font.bold}")
                    print(f"  Tamanho da fonte: {run.font.size}")
                    if run.font.color.rgb:
                        rgb = run.font.color.rgb
                        print(f"  Cor da fonte: RGB({rgb[0]}, {rgb[1]}, {rgb[2]}) -> {run.font.color.rgb}")
                    print(f"  Alinhamento: {paragraph.alignment}")
        
        # Linha 2 - Primeira linha de dados (Meta)
        print("\n" + "=" * 100)
        print("LINHA 2 - DADOS PRIMEIRO TIPO (Meta | 70% | 70% | 70% | 70%)")
        print("=" * 100)
        
        for col_idx, cell in enumerate(tabela.rows[1].cells):
            print(f"\nColuna {col_idx + 1}:")
            print(f"  Texto: {cell.text}")
            
            # Cor de fundo
            tcPr = cell._element.tcPr
            if tcPr is not None:
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    print(f"  Cor de fundo: {fill}")
            
            # Verificar negrito, tamanho e cor da fonte
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    print(f"  Negrito: {run.font.bold}")
                    print(f"  Tamanho da fonte: {run.font.size}")
                    if run.font.color.rgb:
                        rgb = run.font.color.rgb
                        print(f"  Cor da fonte: RGB({rgb[0]}, {rgb[1]}, {rgb[2]})")
                    print(f"  Alinhamento: {paragraph.alignment}")
        
        # Linha 3 - Segunda linha de dados (Resultado)
        print("\n" + "=" * 100)
        print("LINHA 3 - DADOS SEGUNDO TIPO (Resultado | 60.0 | 64.6 | 64.0 | 65.0)")
        print("=" * 100)
        
        for col_idx, cell in enumerate(tabela.rows[2].cells):
            print(f"\nColuna {col_idx + 1}:")
            print(f"  Texto: {cell.text}")
            
            # Cor de fundo
            tcPr = cell._element.tcPr
            if tcPr is not None:
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    print(f"  Cor de fundo: {fill}")
            
            # Verificar negrito, tamanho e cor da fonte
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    print(f"  Negrito: {run.font.bold}")
                    print(f"  Tamanho da fonte: {run.font.size}")
                    if run.font.color.rgb:
                        rgb = run.font.color.rgb
                        print(f"  Cor da fonte: RGB({rgb[0]}, {rgb[1]}, {rgb[2]})")
                    print(f"  Alinhamento: {paragraph.alignment}")
    
    print("\n" + "=" * 100)
    print("✓ VERIFICAÇÃO CONCLUÍDA")
    print("=" * 100)

if __name__ == '__main__':
    verificar_formatacao_detalhada()
