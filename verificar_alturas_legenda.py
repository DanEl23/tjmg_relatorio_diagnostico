#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def verificar_alturas_legenda():
    """Verifica as alturas das linhas e formatação da legenda"""
    doc = Document('teste_metas_institucionais.docx')
    
    print("=" * 100)
    print("VERIFICAÇÃO DE ALTURAS DAS LINHAS E LEGENDA")
    print("=" * 100)
    
    tabelas = doc.tables
    print(f"\n✓ Total de tabelas encontradas: {len(tabelas)}\n")
    
    if len(tabelas) > 0:
        # Analisa primeira tabela
        tabela = tabelas[0]
        print(f"Altura das linhas da primeira tabela (TJMG 5):\n")
        
        for row_idx, row in enumerate(tabela.rows):
            height_emu = row.height
            if height_emu:
                # Converter EMU para cm (1cm = 914400 EMU)
                height_cm = height_emu / 914400
                print(f"  Linha {row_idx + 1}: {height_emu} EMU ≈ {height_cm:.2f}cm")
            else:
                print(f"  Linha {row_idx + 1}: altura não definida")
    
    # Verifica legenda
    print("\n" + "=" * 100)
    print("FORMATAÇÃO DA LEGENDA")
    print("=" * 100 + "\n")
    
    # Procura pela legenda nos parágrafos
    for p_idx, p in enumerate(doc.paragraphs):
        if p.text.startswith('Fonte: Metas'):
            print(f"Parágrafo {p_idx}: {p.text[:60]}...\n")
            print(f"  Estilo: {p.style.name}")
            
            for run in p.runs:
                print(f"  Tamanho da fonte: {run.font.size}")
                print(f"  Itálico: {run.font.italic}")
                if run.font.color.rgb:
                    rgb = run.font.color.rgb
                    print(f"  Cor da fonte: RGB({rgb[0]}, {rgb[1]}, {rgb[2]})")
                else:
                    print(f"  Cor da fonte: padrão ou herdada")
            
            break
    
    print("\n" + "=" * 100)
    print("✓ VERIFICAÇÃO CONCLUÍDA")
    print("=" * 100)

if __name__ == '__main__':
    verificar_alturas_legenda()
