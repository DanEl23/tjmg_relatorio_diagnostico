from docx import Document
from src.tables import builders

def main():
    print("--- 🚀 Teste Específico: Tabela de 4 Colunas ---")
    
    doc = Document()
    
    # DADOS SIMULADOS (Tipo + 4 Colunas de dados)
    # Estrutura esperada: [TIPO, Cód, Descrição, Valor A, Valor B]
    dados_4_cols = [
        ["SUB_HEADER", "AÇÃO ORÇAMENTÁRIA", "CRÉDITO AUTORIZADO 2025 (R$)", "DESPESA REALIZADA 2025 (R$)", "%"],
        ["DATA_ROW", "2053 - Remuneração de Magistrados da Ativa E Encargos Sociais", "1715985406", "1715985406", "100%"],
        ["DATA_ROW", "2054 - Remuneração de Servidores da Ativa e Encargos Sociais", "6275094319", "6275003585", "100%"],
        ["DATA_ROW", "7006 - Proventos de Inativos Civis e Pensionistas", "3119181127", "2954301230", "95%"],
        ["DATA_ROW", "7004 - Precatórios e Sentenças Judiciárias", "1000", "-", "0%"],
        ["TOTAL_ROW", "TOTAL", "11110261852", "10945290221", "99%"]
    ]
    
    doc.add_heading("Teste de Tabela com 4 Colunas", level=1)
    
    builders.adicionar_tabela_orcamento(
        doc, 
        titulo_vindo_do_word="Tabela 4B - Comparativo Anual", 
        dados=dados_4_cols, 
        numero_tabela="04B",
        titulo_custom="Crédito Autorizado x Despesa Realizada por Ação Orçamentária – 2025"
    )
    
    output_file = "teste_4_colunas.docx"
    doc.save(output_file)
    print(f"✅ Arquivo gerado com sucesso: {output_file}")

if __name__ == "__main__":
    main()