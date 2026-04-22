#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def corrigir_marcador_metas():
    """Corrige a posição e ordem do marcador METAS_INSTITUCIONAIS"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("CORRIGINDO MARCADOR: MAPA_RECURSOS::METAS_INSTITUCIONAIS")
    print("=" * 100)
    
    # Procura pelos parágrafos que foram inseridos incorretamente
    indices_para_remover = []
    
    for idx, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if (texto == "MAPA_RECURSOS::METAS_INSTITUCIONAIS" or
            "Apresentação das 55 metas" in texto or
            texto == "11. METAS INSTITUCIONAIS DO TJMG" or
            texto == "[QUEBRA_PAGINA]"):
            indices_para_remover.append(idx)
            print(f"Marcado para remoção [{idx}]: {texto[:50]}")
    
    # Remove os parágrafos em ordem reversa para não afetar os índices
    print(f"\nRemovendo {len(indices_para_remover)} parágrafos...")
    for idx in sorted(indices_para_remover, reverse=True):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
    
    # Procura pela conclusão novamente
    idx_conclusao = None
    for idx, p in enumerate(doc.paragraphs):
        if "CONCLUSÃO" in p.text and "12." in p.text:
            idx_conclusao = idx
            break
    
    if idx_conclusao is None:
        print("❌ Não encontrou CONCLUSÃO")
        return
    
    print(f"\nEncontrando posição de inserção (antes de parágrafo {idx_conclusao})...")
    
    # Agora adiciona corretamente antes da conclusão
    p_insert = doc.paragraphs[idx_conclusao - 1]
    
    # Quebra de página
    p_break = p_insert.insert_paragraph_before("[QUEBRA_PAGINA]")
    
    # Título
    p_titulo = p_insert.insert_paragraph_before("11. METAS INSTITUCIONAIS DO TJMG")
    p_titulo.style = 'Heading 2'
    
    # Descrição
    p_desc = p_insert.insert_paragraph_before(
        "Apresentação das 55 metas institucionais do TJMG para o ano de 2025, com sua evolução "
        "comparada aos anos anteriores (2022-2024) e ataxa de atingimento estabelecida."
    )
    
    # Parágrafo vazio
    p_vazio = p_insert.insert_paragraph_before("")
    
    # Marcador (por último, pois será inserido antes dos anteriores)
    p_marcador = p_insert.insert_paragraph_before("MAPA_RECURSOS::METAS_INSTITUCIONAIS")
    
    # Salva
    doc.save('data/raw/Conteudo_Fonte.docx')
    
    print("✓ Marcador corrigido!")
    print("✓ Documento salvo!")
    print("\nPróximo passo: Execute novamente o gerador com:")
    print("  python main.py --saida Relatorio_Com_Metas_2026.docx")

if __name__ == '__main__':
    corrigir_marcador_metas()
