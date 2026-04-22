from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from .utils import (
    set_cell_vertical_alignment, set_row_height_at_least, 
    set_row_height_flexible, set_cell_bottom_border, set_group_top_border, 
    set_cell_all_borders, remove_all_borders, limpar_espacamento_lista
)

def estilizar_celula(cell, texto, largura, bold, bg_color, align, font_white, remove_bottom=False):
    """Auxiliar para aplicar largura, fundo, bordas pretas e texto."""
    tcPr = cell._element.get_or_add_tcPr()
    
    # Largura
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(largura))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    # Cor de Fundo
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    tcPr.append(shading)
    
    # Bordas (Sempre Pretas)
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{border}')
        # Lógica para remover borda inferior (caso das metas)
        if border == 'bottom' and remove_bottom:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), '000000') # Preto
        tcBorders.append(el)
    tcPr.append(tcBorders)

    # Conteúdo do Texto
    cell.text = texto
    p = cell.paragraphs[0]
    
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left': p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(4)

    run = p.runs[0] if p.runs else p.add_run(texto)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(255, 255, 255) if font_white else RGBColor(0, 0, 0)


def set_vertical_align(cell, align):
    """Auxiliar para centralizar o texto verticalmente na célula."""
    tcPr = cell._element.get_or_add_tcPr()
    tcValign = OxmlElement('w:vAlign')
    tcValign.set(qn('w:val'), align)
    tcPr.append(tcValign)


def aplicar_recuo_tabela(table, recuo_cm):
    """
    Aplica recuo manual em qualquer tabela.
    - table: O objeto tabela criado.
    - recuo_cm: Valor em Centímetros (ex: -0.93 ou 1.5).
    """
    # 1. O alinhamento TEM que ser à Esquerda para o recuo funcionar
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # 2. Converte CM para Twips (1 cm ~ 567 twips)
    valor_twips = int(recuo_cm * 567)
    
    # 3. Acessa as propriedades XML da tabela
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    
    # 4. Cria e aplica a tag de Indentação (w:tblInd)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(valor_twips))
    tblInd.set(qn('w:type'), 'dxa')
    
    # Remove indentação antiga se houver e adiciona a nova
    for child in tblPr.findall(qn('w:tblInd')):
        tblPr.remove(child)
    tblPr.append(tblInd)


def adicionar_tabela_atos(document, dados):
    """ 
    Tabela 01: Atos Normativos 
    Fidelidade Visual: Ajuste fino de espaçamentos (0pt antes, 1.15 entrelinhas no header).
    """
    
    # --- Parâmetros ---
    COR_CABECALHO_HEX = '7F7F7F'                  
    COR_CINZA_CLARO_HEX = 'EEEEEE'                 
    COR_BRANCO_RGB = RGBColor(255, 255, 255)       
    COR_PRETO_RGB = RGBColor(0, 0, 0) 
    
    TAMANHO_FONTE_PADRAO = Pt(12) 
    FONTE = 'Calibri'
    
    # Alturas (Twips)
    ALTURA_HEADER_TWIPS = 397
    ALTURA_DADOS_TWIPS = 227 
    
    # Larguras
    LARGURA_TABELA_TWIPS = '9922' 
    COL_WIDTHS_TWIPS = [2700, 7222]

    # --- Criação ---
    table = document.add_table(rows=1, cols=len(dados[0]))
    
    try: table.style = 'Table Grid'
    except KeyError: pass
        
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # XML: Largura da Tabela
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None: tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), LARGURA_TABELA_TWIPS)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    for i, row_data in enumerate(dados):
        row = table.add_row() if i > 0 else table.rows[0]
        
        # XML: Altura da Linha
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        for existing in trPr.findall(qn('w:trHeight')): trPr.remove(existing)
            
        trHeight = OxmlElement('w:trHeight')
        val_altura = str(ALTURA_HEADER_TWIPS) if i == 0 else str(ALTURA_DADOS_TWIPS)
        trHeight.set(qn('w:val'), val_altura)
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
            
        if i == 0: trPr.append(OxmlElement('w:tblHeader'))
        else:
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0')
            trPr.append(cantSplit)
            
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            
            # XML: Largura da Coluna
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(COL_WIDTHS_TWIPS[j]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
            set_cell_vertical_alignment(cell, 'center') 
            cell.text = "" 
            
            lines = cell_data.split('\n')
            is_first_content_line = True

            for k, line in enumerate(lines):
                line = line.strip()
                if not line: continue 
                
                is_list_item = line.startswith('ü')
                
                if is_first_content_line:
                    current_paragraph = cell.paragraphs[0]
                    is_first_content_line = False
                else:
                    current_paragraph = cell.add_paragraph()

                # --- CORREÇÃO CRÍTICA 1: Espaçamento Antes (Space Before) ---
                # Força ZERO para anular herança do template
                current_paragraph.paragraph_format.space_before = Pt(0)

                text_to_insert = line.replace('ü', '').strip() if is_list_item else line
                run = current_paragraph.add_run(text_to_insert)

                # Formatação Específica por Tipo de Linha
                if is_list_item:
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    estilo_aplicado = False
                    for nome_estilo in ['List Bullet', 'Marcadores', 'Parágrafo com marcadores']:
                        try:
                            current_paragraph.style = nome_estilo
                            estilo_aplicado = True
                            break
                        except KeyError: continue
                    if not estilo_aplicado: run.text = "• " + run.text

                    limpar_espacamento_lista(current_paragraph)
                    
                elif i == 0:
                    # --- CORREÇÃO CRÍTICA 2: Cabeçalho ---
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_paragraph.paragraph_format.space_after = Pt(0)
                    current_paragraph.paragraph_format.line_spacing = 1.15
                else:
                    # Dados Normais
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_paragraph.paragraph_format.space_after = Pt(0)
                    current_paragraph.paragraph_format.line_spacing = 1.0
                
                run.font.name = FONTE
                run.font.size = TAMANHO_FONTE_PADRAO
                run.bold = (i == 0)
                run.font.color.rgb = COR_BRANCO_RGB if i == 0 else COR_PRETO_RGB
            
            # Sombreamento
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), COR_CABECALHO_HEX if i == 0 else (COR_CINZA_CLARO_HEX if i % 2 == 0 else 'auto'))
            if i == 0 or i % 2 == 0:
                 cell._tc.get_or_add_tcPr().append(shading)

    # Legenda
    p_titulo_tabela = document.add_paragraph(style='Normal')
    p_titulo_tabela.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo_tabela.paragraph_format.space_before = Pt(6)
    p_titulo_tabela.paragraph_format.space_after = Pt(12)
    
    run_titulo = p_titulo_tabela.add_run("Tabela 01 - Atos Normativos referentes à Estrutura do TJMG. Fonte: Portal TJMG")
    run_titulo.bold = False 
    run_titulo.font.name = FONTE
    run_titulo.font.size = Pt(8)


def adicionar_tabela_areas(document, dados):
    """ 
    Tabela 02: Principais Áreas 
    Fidelidade Visual: SEM BORDAS (Visual limpo), apenas sombramento e linhas de grupo.
    """
    
    # --- Configurações Físicas ---
    LARGURA_COL1 = 8220  # ~14.5 cm
    LARGURA_COL2 = 1701  # ~3.0 cm
    LARGURA_TOTAL = LARGURA_COL1 + LARGURA_COL2
    ALTURA_LINHA = 227   # 0.4 cm
    
    FONTE_NOME = 'Calibri'
    FONTE_HEADER = Pt(12)
    FONTE_TAM = Pt(11)

    # Cria tabela sem estilo padrão de grade
    table = document.add_table(rows=0, cols=2)
    # Removemos a linha: table.style = 'Table Grid'
    
    # XML: Largura Fixa da Tabela
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None: tblPr.remove(tblW)
        
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- Função Auxiliar Local: Remover Bordas ---
    def limpar_bordas(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil') # 'nil' remove a borda
            tcBorders.append(border)
        
        # Remove definições antigas e aplica a "sem borda"
        old_borders = tcPr.find(qn('w:tcBorders'))
        if old_borders is not None: tcPr.remove(old_borders)
        tcPr.append(tcBorders)

    # --- Função Auxiliar Local: Borda Superior Apenas (Para Grupos) ---
    def adicionar_borda_topo(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        # Topo: Linha sólida preta
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4') # Tamanho 1/2 pt
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        
        # Outros lados: Sem borda
        for side in ['left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)

        old_borders = tcPr.find(qn('w:tcBorders'))
        if old_borders is not None: tcPr.remove(old_borders)
        tcPr.append(tcBorders)

    # --- Função Auxiliar Local: Largura ---
    def set_width(cell, width_twips):
        tcPr = cell._element.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None: tcPr.remove(tcW)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_twips))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    data_row_index = 0 

    for i, row_data in enumerate(dados):
        tipo, col1, col2 = row_data
        row = table.add_row()
        
        # XML: Altura da Linha
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        
        # Remove qualquer altura pré-existente
        for h in trPr.findall(qn('w:trHeight')): 
            trPr.remove(h)
            
        trHeight = OxmlElement('w:trHeight')
        
        # --- Lógica de Altura Diferenciada ---
        if i == 0 or tipo == "HEADER":
            # Altura do Header (ex: 600 twips para ser bem visível)
            trHeight.set(qn('w:val'), '397') 
        else:
            # Altura padrão para as demais linhas
            trHeight.set(qn('w:val'), str(ALTURA_LINHA))
            
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        # Header Repeater
        if tipo.startswith("HEADER"): 
            trPr.append(OxmlElement('w:tblHeader'))
        else: 
            trPr.append(OxmlElement('w:cantSplit'))
        
        c1, c2 = row.cells[0], row.cells[1]

        # 1. Limpa bordas de TUDO inicialmente
        limpar_bordas(c1)
        limpar_bordas(c2)

        # 2. Define larguras iniciais
        set_width(c1, LARGURA_COL1)
        set_width(c2, LARGURA_COL2)

        # --- Lógica de Conteúdo e Estilo ---
        
        # Caso 1: Cabeçalho Principal (DENOMINAÇÃO)
        if tipo == "HEADER_MAIN":
            c1.merge(c2)
            set_width(c1, LARGURA_TOTAL)
            
            # Fundo Cinza Escuro
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '7F7F7F')
            c1._tc.get_or_add_tcPr().append(shading)
            
            c1.text = ""
            p = c1.paragraphs[0]
            run = p.add_run(col1)
            run.font.name = FONTE_NOME
            run.font.size = FONTE_HEADER
            run.font.color.rgb = RGBColor(255, 255, 255) # Texto Branco
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Caso 2: Subtítulos de Grupo (SUPERINTENDÊNCIA...)
        elif tipo == "HEADER_GROUP_SIGLA":
            for j, (cell, texto) in enumerate([(c1, col1), (c2, col2)]):
                # Fundo Cinza Claro
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9D9D9')
                cell._tc.get_or_add_tcPr().append(shading)
                
                # Borda Superior Preta (Importante!)
                adicionar_borda_topo(cell)
                
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(texto)
                run.font.name = FONTE_NOME
                run.font.size = FONTE_TAM
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Texto Preto
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER

        # Caso 3: Subtítulo Mesclado
        elif tipo == "HEADER_GROUP_MERGED":
            c1.merge(c2)
            set_width(c1, LARGURA_TOTAL)
            
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            c1._tc.get_or_add_tcPr().append(shading)
            
            adicionar_borda_topo(c1)
            
            c1.text = ""
            p = c1.paragraphs[0]
            run = p.add_run(col1)
            run.font.name = FONTE_NOME
            run.font.size = FONTE_TAM
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Caso 4: Dados Mesclados
        elif tipo == "DATA_MERGED":
            data_row_index += 1
            c1.merge(c2)
            set_width(c1, LARGURA_TOTAL)
            
            # Zebrado
            if data_row_index % 2 == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'EEEEEE')
                c1._tc.get_or_add_tcPr().append(shading)
                
            c1.text = ""
            p = c1.paragraphs[0]
            run = p.add_run(col1)
            run.font.name = FONTE_NOME
            run.font.size = FONTE_TAM
            run.bold = False
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Caso 5: Dados Divididos (Linhas normais)
        elif tipo == "DATA_SPLIT":
            data_row_index += 1
            cor_fundo = 'EEEEEE' if data_row_index % 2 == 0 else 'auto'
            
            for j, (cell, texto) in enumerate([(c1, col1), (c2, col2)]):
                if cor_fundo != 'auto':
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), cor_fundo)
                    cell._tc.get_or_add_tcPr().append(shading)
                
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(texto)
                run.font.name = FONTE_NOME
                run.font.size = FONTE_TAM
                run.bold = False
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER

        # --- RESET DE ESPAÇAMENTO ---
        for cell in row.cells:
            set_cell_vertical_alignment(cell, 'center')
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

    # Legenda
    p = document.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Tabela 02 - Principais áreas da Secretaria do TJMG. Fonte: Portal TJMG")
    run.font.name = 'Calibri'
    run.font.size = Pt(8)

    
def adicionar_tabela_estrutura(document, dados):
    """ Tabela 03: Estrutura (1 Coluna) - Formatação Rigorosa """
    
    LARGURA_TOTAL = 9911 # ~17.5 cm
    FONTE_HEADER = Pt(12)
    FONTE_TAM = Pt(11)
    ALTURA_LINHA = 227
    
    table = document.add_table(rows=0, cols=1)
    
    aplicar_recuo_tabela(table, -0.5)

    # Configuração de Largura
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    primeiro_grupo_visto = False

    for i, row_data in enumerate(dados):
        tipo, texto = row_data

        # --- Lógica corrigida para o Respiro ---
        if tipo == "HEADER_GROUP_MERGED":
            if primeiro_grupo_visto:
                # Só entra aqui do SEGUNDO grupo em diante
                empty_row = table.add_row()
                trPr_empty = empty_row._tr.get_or_add_trPr()
                trH_empty = OxmlElement('w:trHeight')
                trH_empty.set(qn('w:val'), '15') 
                trH_empty.set(qn('w:hRule'), 'atLeast')
                trPr_empty.append(trH_empty)
                
                # Reset agressivo para a linha não "expandir"
                cell_e = empty_row.cells[0]
                p_e = cell_e.paragraphs[0]
                p_e.paragraph_format.space_before = Pt(0)
                p_e.paragraph_format.space_after = Pt(0)
                p_e.paragraph_format.line_spacing = 0.7
                # Força uma fonte minúscula para garantir altura 15
                run_e = p_e.add_run("")
                run_e.font.size = Pt(1)
                
                remove_all_borders(cell_e)
            else:
                # É o primeiro grupo que encontramos, não fazemos nada e marcamos como visto
                primeiro_grupo_visto = True

        row = table.add_row()
        
        # Altura Manual
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # Remove qualquer altura pré-existente
        for h in trPr.findall(qn('w:trHeight')): 
            trPr.remove(h)
            
        trHeight = OxmlElement('w:trHeight')
        
        # --- Lógica de Altura Diferenciada ---
        if i == 0 or tipo == "HEADER":
            # Altura do Header (ex: 600 twips para ser bem visível)
            trHeight.set(qn('w:val'), '397') 
        else:
            # Altura padrão para as demais linhas
            trHeight.set(qn('w:val'), str(ALTURA_LINHA))

        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        cell = row.cells[0]
        
        # Largura da Célula
        tcPr = cell._element.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(LARGURA_TOTAL))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

        cell.text = texto
        set_cell_vertical_alignment(cell, 'center')

        # Garante que o parágrafo existe e acessamos o run
        p = cell.paragraphs[0]
        if not p.runs:
            p.add_run(texto)
        run = p.runs[0]
        
        if tipo == "HEADER_MAIN":
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '7F7F7F')
            cell._tc.get_or_add_tcPr().append(shading)
            
            p = cell.paragraphs[0]
            p.runs[0].font.color.rgb = RGBColor(255,255,255)
            p.runs[0].bold = True
            p.runs[0].font.size = FONTE_HEADER
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif tipo == "HEADER_GROUP_MERGED":
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            cell._tc.get_or_add_tcPr().append(shading)
            set_group_top_border(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            run.font.size = FONTE_TAM

        else:
            run.font.size = FONTE_TAM

        # --- RESET DE ESPAÇAMENTO ---
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    # Legenda
    p = document.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Tabela 03 - Estruturas para a Prestação Jurisdicional na Segunda Instância. Fonte: Portal TJMG")
    run.font.name = 'Calibri'
    run.font.size = Pt(8)


def adicionar_tabela_comarcas(document, dados):
    """ 
    Tabela 04: Comarcas Instaladas 
    Fidelidade: Dados alinhados à ESQUERDA, Sem Grades, Zebrado.
    """
    
    LARGURA_TOTAL = 9922 
    LARGURA_COL = int(LARGURA_TOTAL / 4) 
    ALTURA_LINHA = 227
    
    FONTE_NOME = 'Calibri'
    FONTE_HEADER = Pt(12)
    FONTE_TAM = Pt(11)

    table = document.add_table(rows=0, cols=4)
    
    # XML: Largura Fixa
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None: tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- Helpers ---
    def limpar_bordas(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil') 
            tcBorders.append(border)
        old_borders = tcPr.find(qn('w:tcBorders'))
        if old_borders is not None: tcPr.remove(old_borders)
        tcPr.append(tcBorders)

    def set_width(cell, width_twips):
        tcPr = cell._element.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None: tcPr.remove(tcW)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_twips))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    data_idx = 0
    
    for i, row_data in enumerate(dados):
        tipo = row_data[0]
        row = table.add_row()
        
        # XML: Altura da Linha
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        for h in trPr.findall(qn('w:trHeight')): 
            trPr.remove(h)
        trHeight = OxmlElement('w:trHeight')
        if i == 0 or tipo == "HEADER":
            trHeight.set(qn('w:val'), '397') 
        else:
            trHeight.set(qn('w:val'), str(ALTURA_LINHA))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        if tipo.startswith("HEADER"):
            trPr.append(OxmlElement('w:tblHeader'))
        else:
            trPr.append(OxmlElement('w:cantSplit'))

        for cell in row.cells:
            limpar_bordas(cell)
            set_width(cell, LARGURA_COL)
            
        # --- Lógica de Conteúdo ---
        
        if tipo == "HEADER_MERGE_4":
            cell = row.cells[0].merge(row.cells[3])
            set_width(cell, LARGURA_TOTAL)
            
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(row_data[1])
            run.font.name = FONTE_NOME
            run.font.size = FONTE_HEADER
            run.font.color.rgb = RGBColor(255,255,255)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '7F7F7F')
            cell._tc.get_or_add_tcPr().append(shading)
            
        elif tipo == "DATA_4_COL":
            data_idx += 1
            cor_fundo = 'D9D9D9' if data_idx % 2 != 0 else 'auto'
            
            for j in range(4):
                cell = row.cells[j]
                
                if cor_fundo != 'auto':
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), cor_fundo)
                    cell._tc.get_or_add_tcPr().append(shading)

                cell.text = ""
                p = cell.paragraphs[0]
                texto_celula = row_data[j+1] if (j+1) < len(row_data) else ""
                
                run = p.add_run(str(texto_celula))
                run.font.name = FONTE_NOME
                run.font.size = FONTE_TAM
                run.font.color.rgb = RGBColor(0,0,0)
                
                # --- CORREÇÃO: ALINHAMENTO À ESQUERDA ---
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # --- RESET GLOBAL DE ESPAÇAMENTO ---
        for cell in row.cells:
            set_cell_vertical_alignment(cell, 'center')
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

    # Legenda
    p = document.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Tabela 04 - Comarcas Instaladas. Fonte: Infoguia")
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
            

def adicionar_tabela_nucleos(document, dados):
    """ 
    Tabela 05: Núcleos (1 Coluna) 
    Fidelidade: Sem grade, Cabeçalhos com borda superior, Espaçamento 0pt.
    """
    
    LARGURA_TOTAL = 9911 # ~17.5 cm
    ALTURA_LINHA = 227
    
    FONTE_NOME = 'Calibri'
    FONTE_TAM = Pt(11)
    
    table = document.add_table(rows=0, cols=1)
    
    aplicar_recuo_tabela(table, -0.5)
    
    # XML: Largura Fixa
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None: tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- Helpers ---
    def limpar_bordas(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        old = tcPr.find(qn('w:tcBorders'))
        if old is not None: tcPr.remove(old)
        tcPr.append(tcBorders)

    def adicionar_borda_topo(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        for side in ['left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        old = tcPr.find(qn('w:tcBorders'))
        if old is not None: tcPr.remove(old)
        tcPr.append(tcBorders)

    for row_data in dados:
        tipo, texto = row_data
        row = table.add_row()
        
        # Altura Manual
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        for h in trPr.findall(qn('w:trHeight')): trPr.remove(h)
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(ALTURA_LINHA))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        cell = row.cells[0]
        limpar_bordas(cell)
        
        # XML: Largura da Célula
        tcPr = cell._element.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None: tcPr.remove(tcW)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(LARGURA_TOTAL))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

        # Lógica de Conteúdo
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(texto)
        run.font.name = FONTE_NOME
        run.font.size = FONTE_TAM
        
        if tipo == "HEADER_GROUP_MERGED":
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            cell._tc.get_or_add_tcPr().append(shading)
            
            adicionar_borda_topo(cell)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        else:
            # Dados normais
            run.bold = False
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Reset Espaçamento
        set_cell_vertical_alignment(cell, 'center')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # Legenda
    p = document.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Tabela 05 - Relação dos Núcleos de Justiça 4.0. Fonte: Infoguia")
    run.font.name = 'Calibri'
    run.font.size = Pt(8)


def adicionar_tabela_processos(document, dados, texto_legenda=None):
    """ 
    Tabelas 06, 07, 08 (Processos, Julgamentos, Acervo)
    Correções: 
    1. Anos: '2022.0' agora vira '2022' (não '20220').
    2. Coluna 1: Texto mantido (não vira '0').
    """
    NUM_COLUNAS = 7
    LARGURA_TOTAL = 9922  # 17.5 cm
    LARGURA_COL_1 = 2600 
    LARGURA_COL_RESTO = int((LARGURA_TOTAL - LARGURA_COL_1) / 6)
    
    H_HEADER = 320 
    ALTURA_LINHA_DADOS = 280 
    
    FONTE_NOME = 'Calibri'
    FONTE_TAM = Pt(12)
    ESPACAMENTO_LINHA = 1.15
    
    COR_TITULO_BG = '7F7F7F'      
    COR_SUB_BG = 'FFFFFF'         
    COR_ZEBRADO = 'D9D9D9'        
    COR_TOTAL_BG = 'D0CECE'       
    COR_DEST_HEADER = '44546A'    
    COR_DEST_DADOS = 'D5DCE4'     
    IDX_DESTAQUE = 5

    # --- FUNÇÕES DE LIMPEZA INTELIGENTE ---
    def converter_para_float(valor):
        """
        Converte valor para float de forma segura.
        1. Tenta conversão direta (para floats do Python: 2022.0 -> 2022.0)
        2. Se falhar, tenta formato BR (remove ponto milhar: 1.234,56 -> 1234.56)
        """
        if isinstance(valor, (int, float)):
            return float(valor)
        
        s = str(valor).strip()
        if not s: return 0.0
        
        try:
            # Tentativa 1: Float padrão (resolve '2022.0')
            return float(s)
        except ValueError:
            try:
                # Tentativa 2: Formato BR (resolve '1.234,56')
                clean_val = s.replace('.', '').replace(',', '.')
                return float(clean_val)
            except:
                return 0.0

    def formatar_brasileiro(valor, is_ano=False):
        if valor is None or str(valor).strip() == "": return ""
        try:
            float_val = converter_para_float(valor)
            
            # Se for ano (2000-2100), exibe sem separador
            if is_ano and (2000 <= float_val <= 2100):
                 return "{:.0f}".format(float_val)
            
            # Formatação BR padrão
            return "{:,.0f}".format(float_val).replace(",", ".")
        except: 
            return str(valor)

    # --- Pre-processamento ---
    novos_dados = []
    totais_colunas = [0.0] * 6 
    tem_total_no_input = any(d[0] == "TOTAL_ROW" for d in dados)

    for d in dados:
        tipo, vals = d[0], list(d[1:])
        if tipo == "DATA_ROW":
            for idx in range(1, 6):
                if idx < len(vals):
                    totais_colunas[idx] += converter_para_float(vals[idx])
            
            nums = [converter_para_float(v) for v in vals[1:6] if str(v).strip()]
            media = sum(nums) / 5 if nums else 0
            
            if len(vals) < 7: vals.append(media)
            else: vals[6] = media
            novos_dados.append([tipo] + vals)
        else:
            novos_dados.append(d)

    if not tem_total_no_input:
        linha_total = ["TOTAL_ROW", "Total"]
        for s in totais_colunas[1:]: linha_total.append(s)
        linha_total.append(sum(totais_colunas[1:]) / 5)
        novos_dados.append(linha_total)

    # --- Construção da Tabela ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    data_row_count = 0
    for row_data in novos_dados:
        tipo, vals = row_data[0], row_data[1:]
        row = table.add_row()
        
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(H_HEADER if tipo in ["HEADER_MERGE", "SUB_HEADER"] else ALTURA_LINHA_DADOS))
        trH.set(qn('w:hRule'), 'atLeast'); trPr.append(trH)

        if tipo == "DATA_ROW": data_row_count += 1

        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            
            tw = OxmlElement('w:tcW')
            tw.set(qn('w:w'), str(LARGURA_COL_1 if j == 0 else LARGURA_COL_RESTO))
            tw.set(qn('w:type'), 'dxa'); tcPr.append(tw)
            
            tcB = OxmlElement('w:tcBorders')
            for s in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{s}'); b.set(qn('w:val'), 'nil'); tcB.append(b)
            tcPr.append(tcB)

            # --- Lógica de Preenchimento ---
            if tipo == "HEADER_MERGE" and j == 0:
                c = cell.merge(row.cells[6])
                c.text = str(vals[0]).upper()
                sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), COR_TITULO_BG); c._tc.get_or_add_tcPr().append(sh)
                p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.runs[0]; r.font.color.rgb = RGBColor(255,255,255); r.bold = True
            
            elif tipo == "SUB_HEADER":
                # --- CORREÇÃO: Coluna 0 é texto puro, Colunas 1-5 são anos ---
                if j == 0:
                    texto = str(vals[j]) # Mantém "Instância"
                elif j == 6:
                    texto = "Média"
                else:
                    texto = formatar_brasileiro(vals[j], is_ano=True)
                
                cell.text = texto
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.runs[0]; r.bold = True
                
                tcB_sub = cell._element.get_or_add_tcPr().find(qn('w:tcBorders'))
                bt = OxmlElement('w:bottom'); bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '6'); bt.set(qn('w:color'), '000000')
                tcB_sub.append(bt)
                
                bg = '44546A' if j == IDX_DESTAQUE else COR_SUB_BG
                if j == IDX_DESTAQUE: r.font.color.rgb = RGBColor(255,255,255)
                sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), bg); cell._tc.get_or_add_tcPr().append(sh)

            elif tipo in ["DATA_ROW", "TOTAL_ROW"]:
                is_total = (tipo == "TOTAL_ROW")
                # Coluna 0 sempre texto, demais formata
                cell.text = formatar_brasileiro(vals[j]) if j > 0 else str(vals[j])
                
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                r = p.runs[0]; r.bold = is_total
                
                if j == IDX_DESTAQUE: bg = COR_DEST_DADOS
                elif is_total: bg = COR_TOTAL_BG
                elif data_row_count % 2 == 0: bg = COR_ZEBRADO
                else: bg = 'FFFFFF'
                sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), bg); cell._tc.get_or_add_tcPr().append(sh)

            set_cell_vertical_alignment(cell, 'center')
            for p in cell.paragraphs:
                p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
                for r in p.runs: 
                    r.font.name = FONTE_NOME; r.font.size = FONTE_TAM
                    if tipo == "TOTAL_ROW": r.bold = True

    # --- Legenda ---
    p_leg = document.add_paragraph()
    p_leg.paragraph_format.space_before = Pt(6)
    
    if texto_legenda:
            # Se o Word já mandou o título completo (ex: "Tabela 06 - ..."), usamos ele direto
            texto_final = texto_legenda.strip()
    else:
        # Fallback de segurança (caso não venha do Word)
        titulo_dados = str(novos_dados[0][1]).upper()
        prefixo = "08"
        if "PROCESSOS" in titulo_dados: prefixo = "06"
        elif "JULGAMENTOS" in titulo_dados: prefixo = "07"
        texto_final = f"Tabela {prefixo} - {novos_dados[0][1]}"

    # Monta a legenda final sem duplicar o prefixo
    txt_leg = f"{texto_final.rstrip('.')}. Fonte: Centro de Informações para a Gestão Institucional – CEINFO."    
    r_leg = p_leg.add_run(txt_leg); r_leg.font.name = FONTE_NOME; r_leg.font.size = Pt(8)

    
def adicionar_tabela_orcamento(document, titulo_vindo_do_word, dados, numero_tabela="09", titulo_custom=None):
    """
    Tabela de Orçamento Dinâmica.
    - Atualização: Alinhamento ESQUERDA com Recuo Manual (tblInd) para ajuste fino de centralização.
    """
    if not dados: return

    # --- 1. DETECÇÃO DE COLUNAS ---
    sample_row = next((d for d in dados if d[0] == 'DATA_ROW'), dados[0])
    qtd_dados = len(sample_row) - 1 
    NUM_COLUNAS = qtd_dados 
    
    # --- Configurações Padrão ---
    LARGURA_TOTAL = 9922 
    ALTURA_LINHA = 340 
    FONTE_NOME = 'Calibri'; FONTE_TAM = Pt(12); ESPACAMENTO_LINHA = 1.15
    COR_HEADER_BG = '7F7F7F'; COR_TOTAL_BG = 'BFBFBF'; COR_DADOS_BG = 'FFFFFF'
    
    # Variável de Recuo (Default 0)
    RECUO_TABELA = 0 

    # --- DEFINIÇÃO DE LARGURAS ---
    if NUM_COLUNAS == 4:
        # Layout 4 Colunas (Valores Fixos)
        LARGURA_COL_1 = 4184
        LARGURA_COL_2 = 2680
        LARGURA_COL_3 = 2680
        LARGURA_COL_4 = 793
        larguras = [LARGURA_COL_1, LARGURA_COL_2, LARGURA_COL_3, LARGURA_COL_4]
        
        LARGURA_TOTAL_REAL = sum(larguras) # ~10337 twips (18.2cm)
        ALTURA_LINHA = 646 # 1.14cm
        
        # --- AJUSTE DE RECUO AQUI ---
        # Como a tabela tem 18.2cm e a área útil é ~16cm, ela sobra 2.2cm.
        # Para centralizar visualmente, precisamos puxar para a esquerda (negativo).
        # Tente valores como -500, -900, -1200 até achar o centro ideal.
        # Se quiser empurrar para a DIREITA, use valor positivo.
        RECUO_TABELA = -675  

    elif NUM_COLUNAS == 3:
        LARGURA_COL_1 = 1500; LARGURA_COL_2 = 5422; LARGURA_COL_3 = 3000  
        larguras = [LARGURA_COL_1, LARGURA_COL_2, LARGURA_COL_3]
        LARGURA_TOTAL_REAL = 9922
    else:
        LARGURA_COL_1 = 6500; LARGURA_COL_2 = 9922 - LARGURA_COL_1
        larguras = [LARGURA_COL_1, LARGURA_COL_2]
        LARGURA_TOTAL_REAL = 9922

    # --- 2. TÍTULO SUPERIOR ---
    if titulo_custom:
        p_top = document.add_paragraph()
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_top = p_top.add_run(str(titulo_custom))
        run_top.font.name = FONTE_NOME; run_top.font.size = Pt(12); run_top.bold = True
        p_top.paragraph_format.space_after = Pt(12)

    # --- 3. CONSTRUÇÃO DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.autofit = False 
    
    # [IMPORTANTE] Mudamos para LEFT para o recuo (tblInd) funcionar corretamente
    table.alignment = WD_TABLE_ALIGNMENT.LEFT 

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Define Largura Total Fixa
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL)); tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)
    
    # Trava Layout
    tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)

    # [NOVO] Aplica o Recuo (Indentation)
    if RECUO_TABELA != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(RECUO_TABELA))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # Helpers
    def converter_para_float(valor):
        if isinstance(valor, (int, float)): return float(valor)
        s = str(valor).strip().replace('R$', '').replace(' ', '')
        if not s or s == '-': return 0.0
        try: return float(s.replace('.', '').replace(',', '.'))
        except: return 0.0

    def formatar_moeda(valor):
        val_float = converter_para_float(valor)
        if val_float == 0: return "-"
        return "{:,.2f}".format(val_float).replace(",", "X").replace(".", ",").replace("X", ".")

    # --- LOOP DE DADOS ---
    for i, row_data in enumerate(dados):
        tipo = row_data[0]
        vals = [str(x) for x in row_data[1:]]
        while len(vals) < NUM_COLUNAS: vals.append("")

        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(ALTURA_LINHA))
        trH.set(qn('w:hRule'), 'atLeast') 
        trPr.append(trH)

        if i == 0 or tipo == "SUB_HEADER": trPr.append(OxmlElement('w:tblHeader'))

        if tipo == "GROUP_TITLE":
            cell = row.cells[0]; cell.merge(row.cells[NUM_COLUNAS - 1])
            cell.text = vals[0].upper()
            sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), COR_HEADER_BG); cell._tc.get_or_add_tcPr().append(sh)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.runs[0]; r.font.name = FONTE_NOME; r.font.size = FONTE_TAM; r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
            tcPr = cell._element.get_or_add_tcPr(); tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{edge}'); border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:color'), '000000'); tcBorders.append(border)
            tcPr.append(tcBorders)
            continue 

        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            largura_atual = larguras[j] if j < len(larguras) else int(LARGURA_TOTAL_REAL / NUM_COLUNAS)
            tw = OxmlElement('w:tcW'); tw.set(qn('w:w'), str(largura_atual)); tw.set(qn('w:type'), 'dxa'); tcPr.append(tw)
            
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{edge}'); border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:color'), '000000'); tcBorders.append(border)
            tcPr.append(tcBorders)

            texto_f = vals[j]
            if tipo == "SUB_HEADER": texto_f = texto_f.upper()
            
            if tipo in ["DATA_ROW", "TOTAL_ROW"]:
                if NUM_COLUNAS == 2 and j == 1: texto_f = formatar_moeda(vals[j])
                elif NUM_COLUNAS == 3 and j > 0:
                     if any(c.isdigit() for c in vals[j]): texto_f = formatar_moeda(vals[j])
                elif NUM_COLUNAS == 4 and j in [1, 2]: texto_f = formatar_moeda(vals[j])

            cell.text = texto_f
            p = cell.paragraphs[0]; r = p.runs[0] if p.runs else p.add_run(texto_f)
            
            bg = COR_DADOS_BG; is_bold = False; f_color = '000000'
            if tipo == "SUB_HEADER": bg = COR_HEADER_BG; is_bold = True; f_color = 'FFFFFF'
            elif tipo == "TOTAL_ROW": bg = COR_TOTAL_BG; is_bold = True
            
            sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), bg); cell._tc.get_or_add_tcPr().append(sh)
            set_cell_vertical_alignment(cell, 'center')
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
            
            if j == 0 and tipo != "SUB_HEADER":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.left_indent = Pt(6)
            else: 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            r.font.name = FONTE_NOME; r.font.size = FONTE_TAM; r.bold = is_bold; r.font.color.rgb = RGBColor.from_string(f_color)

    # --- 4. LEGENDA INFERIOR ---
    p_leg = document.add_paragraph()
    p_leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg.paragraph_format.space_before = Pt(6)
    fonte_final = "Armazém de Informações - BO SIAFI/MG" if str(numero_tabela) in ["09", "10"] else "LOA"
    titulo_formatado = titulo_vindo_do_word.strip()
    if not titulo_formatado.upper().startswith("TABELA"):
        titulo_formatado = f"Tabela {numero_tabela} - {titulo_formatado}"    
    texto_legenda = f"{titulo_formatado}. Fonte: {fonte_final}"
    r_leg = p_leg.add_run(texto_legenda)
    r_leg.font.name = FONTE_NOME; r_leg.font.size = Pt(8)


def adicionar_tabela_orcamento_conjunto(document, dados):
    """
    Tabela 11: Orçamento Conjunto.
    - Atualização: Inclui 'R$' na formatação da Coluna 2.
    - Ex: R$ 1.234.567
    """
    if not dados: return

    # --- Configurações Físicas ---
    LARGURA_TOTAL = 9922 
    LARGURA_COL_2 = 4111 
    LARGURA_COL_1 = LARGURA_TOTAL - LARGURA_COL_2
    
    FONTE_NOME = 'Calibri'
    TAMANHO_FONTE = Pt(12)
    ALTURA_LINHA_TWIPS = '340' 
    COR_GROUP_BG = '7F7F7F'    
    COR_SUB_BG = 'D9D9D9'      
    COR_TOTAL_BG = 'BFBFBF'    
    
    # Cria a Tabela
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(LARGURA_TOTAL)); tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)

    # --- HELPER: Formatação de Moeda (R$ + Inteiro + Milhar) ---
    def formatar_moeda(valor):
        if not valor: return ""
        try:
            # Limpeza e conversão
            if isinstance(valor, str):
                s = valor.replace('R$', '').replace(' ', '').strip()
                if not s or s in ['-', 'nan', 'None']: return "-"
                val_float = float(s.replace('.', '').replace(',', '.'))
            else:
                val_float = float(valor)
            
            # 1. Formata o número: 1.234.567
            numero_str = "{:,.0f}".format(val_float).replace(",", ".")
            
            # 2. Adiciona o prefixo R$
            return f"R$ {numero_str}"
        except:
            return str(valor)

    # --- HELPER: Formatação de Célula ---
    def formatar_celula(cell, texto, bold=False, color_rgb=RGBColor(0,0,0), align_h='LEFT', bg_color=None, largura=None):
        cell.text = str(texto) if texto else ""
        p = cell.paragraphs[0]
        
        if align_h == 'CENTER': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align_h == 'RIGHT': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        set_cell_vertical_alignment(cell, 'center')
        
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(0)

        run = p.runs[0] if p.runs else p.add_run(str(texto) if texto else "")
        run.font.name = FONTE_NOME
        run.font.size = TAMANHO_FONTE
        run.bold = bold
        run.font.color.rgb = color_rgb
        
        if bg_color:
            tcPr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), bg_color); tcPr.append(shd)
        
        if largura:
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(largura)); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)

        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{edge}'); border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), '000000'); tcBorders.append(border)
        tcPr.append(tcBorders)

    def definir_altura_linha(row):
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight'); trH.set(qn('w:val'), ALTURA_LINHA_TWIPS); trH.set(qn('w:hRule'), 'atLeast'); trPr.append(trH)

    # ==========================================
    # LOOP PRINCIPAL
    # ==========================================
    for row_data in dados:
        tipo = row_data[0]
        vals = row_data[1:] 
        
        val1 = vals[0] if len(vals) > 0 else ""
        val2_raw = vals[1] if len(vals) > 1 else ""

        row = table.add_row()
        definir_altura_linha(row)

        if tipo == "GROUP_TITLE":
            cell = row.cells[0]
            cell.merge(row.cells[1])
            formatar_celula(cell, val1.upper(), bold=True, color_rgb=RGBColor(255, 255, 255), align_h='CENTER', bg_color=COR_GROUP_BG, largura=LARGURA_TOTAL)

        elif tipo == "SUB_HEADER":
            formatar_celula(row.cells[0], val1, bold=True, align_h='CENTER', bg_color=COR_SUB_BG, largura=LARGURA_COL_1)
            formatar_celula(row.cells[1], val2_raw, bold=True, align_h='CENTER', bg_color=COR_SUB_BG, largura=LARGURA_COL_2)

        elif tipo == "TOTAL_ROW":
            # Formata moeda na coluna 2
            val2_fmt = formatar_moeda(val2_raw)
            
            formatar_celula(row.cells[0], val1, bold=True, align_h='CENTER', bg_color=COR_TOTAL_BG, largura=LARGURA_COL_1)
            formatar_celula(row.cells[1], val2_fmt, bold=True, align_h='CENTER', bg_color=COR_TOTAL_BG, largura=LARGURA_COL_2)

        else:
            # Formata moeda na coluna 2
            val2_fmt = formatar_moeda(val2_raw)

            formatar_celula(row.cells[0], val1, bold=False, align_h='LEFT', largura=LARGURA_COL_1)
            formatar_celula(row.cells[1], val2_fmt, bold=False, align_h='CENTER', largura=LARGURA_COL_2)

    # --- LEGENDA ---
    p_leg = document.add_paragraph()
    p_leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg.paragraph_format.space_before = Pt(6)
    r_leg = p_leg.add_run("Tabela 11 - Orçamento 2026 por ação orçamentária. Fonte: Lei Orçamentária Anual nº 25.698, de 14/01/2026.")
    r_leg.font.name = FONTE_NOME
    r_leg.font.size = Pt(8)


def adicionar_tabela_orcamento_detalhada(document, dados):
    """
    Tabela 4 Colunas com LARGURAS e RECUO FIXOS.
    - Larguras: [6000, 2000, 800, 2000]
    - Recuo: -0,93 cm (-527 twips)
    """
    if not dados: return

    # ==========================================================================
    # 1. CONFIGURAÇÕES FIXAS (Hardcoded)
    # ==========================================================================
    # Larguras exatas solicitadas
    LARGURAS = [5500, 2000, 950, 2000] 
    
    # Soma total (10800 twips)
    LARGURA_TOTAL_REAL = sum(LARGURAS)
    
    # Configurações Visuais
    FONTE_NOME = 'Calibri'; TAMANHO_FONTE = Pt(11); ALTURA_LINHA_TWIPS = '340'
    COR_GROUP_BG = '7F7F7F'; COR_SUB_BG = 'D9D9D9'; COR_TOTAL_BG = 'BFBFBF'

    # ==========================================================================
    # 2. HELPERS
    # ==========================================================================
    def formatar_moeda(valor):
        if not valor: return ""
        try:
            if isinstance(valor, str):
                s = valor.replace('R$', '').replace(' ', '').strip()
                if not s or s in ['-', 'nan', 'None']: return "-"
                val_float = float(s.replace('.', '').replace(',', '.'))
            else: val_float = float(valor)
            return f"R$ " + "{:,.0f}".format(val_float).replace(",", ".")
        except: return str(valor)

    def formatar_celula(cell, texto, bold=False, bg_color=None, align='CENTER', largura=None):
        cell.text = str(texto) if texto else ""
        p = cell.paragraphs[0]
        
        # Alinhamento
        if align == 'LEFT': p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'RIGHT': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        set_cell_vertical_alignment(cell, 'center')
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        if align == 'LEFT': p.paragraph_format.left_indent = Pt(6)

        run = p.runs[0] if p.runs else p.add_run(str(texto) if texto else "")
        run.font.name = FONTE_NOME; run.font.size = TAMANHO_FONTE; run.bold = bold
        
        if bg_color == COR_GROUP_BG: run.font.color.rgb = RGBColor(255, 255, 255)
        else: run.font.color.rgb = RGBColor(0, 0, 0)
        
        if bg_color:
            tcPr = cell._element.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), bg_color); tcPr.append(shd)
        
        # Aplica Largura Fixa na Célula
        if largura:
            tcPr = cell._element.get_or_add_tcPr(); tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(largura)); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
        
        # Bordas
        tcPr = cell._element.get_or_add_tcPr(); tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{edge}'); border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), '000000'); tcBorders.append(border)
        tcPr.append(tcBorders)

    # ==========================================================================
    # 3. CRIAÇÃO DA TABELA
    # ==========================================================================
    table = document.add_table(rows=0, cols=4)
    table.autofit = False 
    table.alignment = WD_TABLE_ALIGNMENT.CENTER # Obrigatório para o recuo funcionar
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Define Largura Total e Layout Fixo
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL)); tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)
    
    # ==========================================================================
    # 4. PREENCHIMENTO
    # ==========================================================================
    for row_data in dados:
        tipo = row_data[0]
        vals = list(row_data[1:])
        while len(vals) < 4: vals.append("")

        row = table.add_row()
        trPr = row._tr.get_or_add_trPr(); trH = OxmlElement('w:trHeight'); trH.set(qn('w:val'), ALTURA_LINHA_TWIPS); trH.set(qn('w:hRule'), 'atLeast'); trPr.append(trH)

        # Caso Group Title (Mesclado)
        if tipo == "GROUP_TITLE":
            cell = row.cells[0]; cell.merge(row.cells[3])
            formatar_celula(cell, vals[0].upper(), bold=True, bg_color=COR_GROUP_BG, align='CENTER', largura=LARGURA_TOTAL_REAL)
            continue

        bg_atual = None; bold_atual = False
        if tipo == "SUB_HEADER": bg_atual = COR_SUB_BG; bold_atual = True
        elif tipo == "TOTAL_ROW": bg_atual = COR_TOTAL_BG; bold_atual = True

        for j, cell in enumerate(row.cells):
            valor = vals[j]
            
            # Coluna 1 (Texto)
            if j == 0:
                align_atual = 'LEFT' if tipo != "SUB_HEADER" else 'CENTER'
                valor_final = valor.upper() if tipo == "SUB_HEADER" else valor
            
            # Colunas 2, 3, 4 (Valores)
            else:
                align_atual = 'CENTER'
                if tipo in ["DATA_ROW", "TOTAL_ROW"]: valor_final = formatar_moeda(valor)
                else: valor_final = valor.upper() 

            # Usa a largura fixa da lista LARGURAS
            formatar_celula(cell, valor_final, bold=bold_atual, bg_color=bg_atual, align=align_atual, largura=LARGURAS[j])

    p_leg = document.add_paragraph()
    p_leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg.paragraph_format.space_before = Pt(6)
    r_leg = p_leg.add_run("Tabela Orçamentária Detalhada. Fonte: LOA/TJMG")
    r_leg.font.name = FONTE_NOME; r_leg.font.size = Pt(8)
    

def adicionar_tabela_cidades(document, dados):
    """ Tabela 12: Cidades (4 colunas zebradas) """
    table = document.add_table(rows=0, cols=4)
    tbl = table._tbl
    
    # CORREÇÃO DO WARNING AQUI
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
        
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Cm(1.27).twips)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    data_idx = 0
    for row_data in dados:
        tipo = row_data[0]
        row = table.add_row()
        set_row_height_flexible(row, 584)
        if tipo == "DATA_ROW": data_idx += 1
        
        for j in range(4):
            cell = row.cells[j]
            cell.text = row_data[j+1]
            set_cell_all_borders(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if tipo == "DATA_ROW" and data_idx % 2 != 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9D9D9')
                cell._tc.get_or_add_tcPr().append(shading)


def adicionar_tabela_justica_numeros(document, dados, texto_legenda=None, indent_cm=-0.5, fonte=None):
    """ 
    Tabela 13: Justiça em Números 
    - Estrutura: XML Fixed (Mantida intacta)
    - Legenda: Rodapé unificado (Título + Fonte) estilo Normal/Calibri 8pt.
    """
    if not dados: return

    # --- 1. CRIAÇÃO DA ESTRUTURA DA TABELA (CÓDIGO XML INTACTO) ---
    table = document.add_table(rows=0, cols=7)
    tbl = table._tbl
    
    # Configuração da Tabela
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
        
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Recuo
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips))) 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Larguras
    tblGrid = OxmlElement('w:tblGrid')
    widths = [Cm(5.5)] + [Cm(2.25)]*6
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w.twips)))
        tblGrid.append(gc)
    tbl.insert(1, tblGrid)
    
    # Preenchimento
    data_idx = 0
    for row_data in dados:
        tipo = row_data[0]
        vals = row_data[1:]
        row = table.add_row()
        set_row_height_flexible(row, 272)
        
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        if tipo.startswith("HEADER") or tipo.startswith("SUB"):
            trPr.append(OxmlElement('w:tblHeader'))
            
        # Header Mesclado
        if tipo == "HEADER_MERGE":
            c = row.cells[0].merge(row.cells[6])
            c.text = vals[0]
            set_cell_vertical_alignment(c, 'center') 
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '44546A')
            c._tc.get_or_add_tcPr().append(shading)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Calibri' 
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255,255,255)
                run.bold = True
            remove_all_borders(c)

        # Subtítulos
        elif tipo in ["SUB_HEADER", "SUB_HEADER_SECONDARY"]:
            for j in range(7):
                c = row.cells[j]
                c.text = vals[j]
                set_cell_vertical_alignment(c, 'center') 
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'EEEEEE')
                c._tc.get_or_add_tcPr().append(shading)
                remove_all_borders(c)
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                if p.runs:
                    run = p.runs[0]
                    run.font.name = 'Calibri' 
                    run.font.size = Pt(11)
                    run.bold = True
                if tipo == "SUB_HEADER_SECONDARY":
                    set_cell_bottom_border(c)
                    
        # Dados
        elif tipo == "DATA_ROW":
            data_idx += 1
            for j in range(7):
                c = row.cells[j]
                c.text = vals[j]
                set_cell_vertical_alignment(c, 'center') 
                remove_all_borders(c)
                if data_idx % 2 != 0:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D9D9D9')
                    c._tc.get_or_add_tcPr().append(shading)
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                if p.runs:
                    run = p.runs[0]
                    run.font.name = 'Calibri' 
                    run.font.size = Pt(11)

    # --- 2. AQUI ESTÁ A LÓGICA DA LEGENDA (IGUAL TABELA ESTRUTURA) ---
    # Unifica Título + Fonte em um único parágrafo pequeno abaixo da tabela.
    
    if texto_legenda or fonte:
        p = document.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        
        # Define um texto padrão se a fonte não vier preenchida
        texto_fonte = fonte if fonte else "Fonte: Base de Dados Justiça em Números."
        
        # Concatenação Inteligente: Título + Ponto + Fonte
        if texto_legenda:
            # Garante que não fique ponto duplo (.. ou . .)
            titulo_limpo = texto_legenda.strip()
            if titulo_limpo.endswith('.'):
                titulo_limpo = titulo_limpo[:-1]
                
            texto_final = f"{titulo_limpo}. {texto_fonte}"
        else:
            texto_final = texto_fonte
            
        run = p.add_run(texto_final)
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        

def adicionar_tabela_generica(document, titulo_tabela, dados, fonte=None):
    """
    Tabela Genérica (2 colunas) - Layout Ajustado Final.
    - Aceita fonte customizada.
    - Mantém tratamento de marcadores 'ü'.
    """
    if not dados: return

    # --- Configurações Físicas ---
    LARGURA_TOTAL = 9922 
    LARGURA_COL_1 = int(LARGURA_TOTAL * 0.3)
    LARGURA_COL_2 = LARGURA_TOTAL - LARGURA_COL_1
    
    FONTE_NOME = 'Calibri'
    ESPACAMENTO_LINHA = 1.0  
    COR_HEADER_BG = '7F7F7F'

    # --- Tabela ---
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(LARGURA_TOTAL)); tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)

    for i, linha in enumerate(dados):
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), '340')
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)

        texto_col1 = str(linha[0])
        texto_col2 = str(linha[1])

        # --- CABEÇALHO ---
        eh_cabecalho = (i == 0)
        if eh_cabecalho:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
            texto_col1 = texto_col1.upper()
            texto_col2 = texto_col2.upper()

        vals = [texto_col1, texto_col2]
        
        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            
            # Larguras
            tw = OxmlElement('w:tcW')
            tw.set(qn('w:w'), str(LARGURA_COL_1 if j == 0 else LARGURA_COL_2))
            tw.set(qn('w:type'), 'dxa')
            tcPr.append(tw)
            
            # Bordas
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            
            set_cell_vertical_alignment(cell, 'center')

            # --- CONTEÚDO ---
            texto_base = vals[j]
            
            # CASO ESPECIAL: Marcador 'ü' na área de dados
            if not eh_cabecalho and "ü" in texto_base:
                cell.text = "" # Limpa conteúdo
                
                primeiro_uso = True
                itens = texto_base.split('ü')
                
                for k, item in enumerate(itens):
                    item_limpo = item.strip()
                    if not item_limpo: continue 
                    
                    if primeiro_uso:
                        p_atual = cell.paragraphs[0]
                        primeiro_uso = False
                    else:
                        p_atual = cell.add_paragraph()
                    
                    p_atual.text = item_limpo
                    
                    # Lógica de Estilo (Intro vs Bullet)
                    eh_intro = (k == 0 and not texto_base.strip().startswith("ü"))
                    
                    if eh_intro:
                         p_atual.style = None 
                         p_atual.paragraph_format.left_indent = Pt(0) 
                    else:
                         p_atual.style = 'List Bullet'
                    
                    p_atual.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_atual.paragraph_format.line_spacing = ESPACAMENTO_LINHA
                    p_atual.paragraph_format.space_before = Pt(0)
                    p_atual.paragraph_format.space_after = Pt(0)
                    
                    if p_atual.runs:
                        run = p_atual.runs[0]
                        run.font.name = FONTE_NOME
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0,0,0)
                        run.bold = False

            else:
                # --- TEXTO NORMAL ---
                cell.text = texto_base
                p = cell.paragraphs[0]
                
                p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Pt(0)

                run = p.runs[0] if p.runs else p.add_run(texto_base)
                run.font.name = FONTE_NOME
                run.font.size = Pt(12)

                if eh_cabecalho:
                    sh = OxmlElement('w:shd')
                    sh.set(qn('w:fill'), COR_HEADER_BG)
                    cell._tc.get_or_add_tcPr().append(sh)
                    
                    run.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    run.bold = False
                    run.font.color.rgb = RGBColor(0,0,0)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.LEFT

    # --- LEGENDA (ALTERADA AQUI) ---
    p_leg = document.add_paragraph()
    p_leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg.paragraph_format.space_before = Pt(6)
    
    # Lógica de Seleção da Fonte
    if fonte:
        fonte_texto = fonte
    else:
        # Fallback original
        fonte_texto = "CEINFO" if "Atos" in titulo_tabela else "TJMG"

    titulo_formatado = titulo_tabela.strip()
    
    # Garante que termine com ponto antes de "Fonte:"
    if titulo_formatado.endswith('.'):
        titulo_formatado = titulo_formatado[:-1]
        
    texto_legenda = f"{titulo_formatado}. Fonte: {fonte_texto}"
    
    r_leg = p_leg.add_run(texto_legenda)
    r_leg.font.name = FONTE_NOME
    r_leg.font.size = Pt(8)


def adicionar_tabela_simples_3col(document, dados, titulo_custom=None, indent_cm=0, fonte=None):
    """
    Tabela de 3 colunas com legenda unificada no rodapé (Título + Fonte).
    """
    if not dados: return

    # --- 1. CONFIGURAÇÕES ---
    NUM_COLUNAS = 3
    LARGURA_COL_1 = 2000  
    LARGURA_COL_2 = 2500  
    LARGURA_COL_3 = 4422  
    larguras = [LARGURA_COL_1, LARGURA_COL_2, LARGURA_COL_3]
    LARGURA_TOTAL_REAL = sum(larguras)
    
    ALTURA_LINHA = 340 
    FONTE_NOME = 'Calibri'
    FONTE_TAM = Pt(11)
    ESPACAMENTO_LINHA = 1.15
    
    COR_HEADER_BG = '44546A'
    COR_SUBHEADER_BG = 'D9D9D9'
    COR_DADOS_BG_PAR = 'F2F2F2' 
    COR_DADOS_BG_IMPAR = 'D9D9D9'

    # --- (REMOVIDO: TÍTULO SUPERIOR) ---
    # O título agora será inserido apenas no final.

    # --- 2. ESTRUTURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.autofit = False 
    table.alignment = WD_TABLE_ALIGNMENT.LEFT 

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Aplica Recuo
    if indent_cm != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # --- 3. POPULAÇÃO DOS DADOS ---
    data_idx = 0
    for i, row_data in enumerate(dados):
        tipo = row_data[0] 
        vals = [str(x) for x in row_data[1:]]
        while len(vals) < NUM_COLUNAS: vals.append("")
        
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(ALTURA_LINHA))
        trH.set(qn('w:hRule'), 'atLeast') 
        trPr.append(trH)

        if tipo in ["HEADER", "SUB_HEADER"]:
            trPr.append(OxmlElement('w:tblHeader'))
        else:
            trPr.append(OxmlElement('w:cantSplit'))

        # Header Mesclado (Barra Azul Interna)
        if tipo == "HEADER":
            c = row.cells[0].merge(row.cells[NUM_COLUNAS-1])
            c.text = vals[0].upper()
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            tcPr = c._element.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), COR_HEADER_BG)
            tcPr.append(shading)
            set_cell_vertical_alignment(c, 'center')
            continue 

        # Colunas de Dados
        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(larguras[j]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

            cell.text = vals[j]
            p = cell.paragraphs[0]
            
            if j == 1: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = FONTE_TAM
            run.font.color.rgb = RGBColor(0, 0, 0)

            shading = OxmlElement('w:shd')
            if tipo == "SUB_HEADER":
                run.bold = True
                shading.set(qn('w:fill'), COR_SUBHEADER_BG)
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)
            else: 
                run.bold = False
                bg_color = COR_DADOS_BG_PAR if data_idx % 2 == 0 else COR_DADOS_BG_IMPAR
                shading.set(qn('w:fill'), bg_color)
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), 'D9D9D9')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)
            
            cell._tc.get_or_add_tcPr().append(shading)
            set_cell_vertical_alignment(cell, 'center')
            p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        
        if tipo == "DATA_ROW": data_idx += 1

    # --- 4. LEGENDA INFERIOR (TÍTULO + FONTE EM PRETO) ---
    if titulo_custom or fonte:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        
        # Constrói o texto unificado
        texto_completo = ""
        if titulo_custom:
            texto_completo += str(titulo_custom).strip().rstrip('.') + ". "
        
        # Adiciona a fonte se existir, senão usa um padrão
        txt_fonte = fonte if fonte else "Fonte: TJMG."
        texto_completo += txt_fonte
        
        run = p.add_run(texto_completo)
        run.font.name = 'Calibri'
        run.font.size = Pt(9) # Tamanho legível para legenda
        run.font.color.rgb = RGBColor(0, 0, 0) # PRETO


def adicionar_tabela_4col_simples(document, dados, titulo_custom=None, indent_cm=0, fonte=None, larguras=None):
    """
    Tabela de 4 colunas.
    Aceita 'larguras' (lista de 4 inteiros em twips) para personalização.
    Caso contrário, usa o padrão (1 larga + 3 iguais).
    """
    if not dados: return

    # --- 1. CONFIGURAÇÕES DE LARGURA ---
    NUM_COLUNAS = 4
    
    # Se o usuário passou larguras personalizadas, usa elas.
    # Senão, usa o padrão (Col 1 Larga, Resto Igual).
    if larguras and len(larguras) == NUM_COLUNAS:
        lista_larguras = larguras
    else:
        # Padrão Default
        LARGURA_COL_1 = 3000  # ~5.3 cm
        LARGURA_RESTANTE = 2212 # ~3.9 cm
        lista_larguras = [LARGURA_COL_1, LARGURA_RESTANTE, LARGURA_RESTANTE, LARGURA_RESTANTE]

    LARGURA_TOTAL_REAL = sum(lista_larguras)
    
    ALTURA_LINHA = 340 
    FONTE_NOME = 'Calibri'
    FONTE_TAM = Pt(11)
    ESPACAMENTO_LINHA = 1.15
    
    # Cores
    COR_HEADER_BG = '44546A'    
    COR_SUBHEADER_BG = 'D9D9D9' 
    COR_DADOS_BG_PAR = 'F2F2F2' 
    COR_DADOS_BG_IMPAR = 'D9D9D9' 

    # --- 2. ESTRUTURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.autofit = False 
    table.alignment = WD_TABLE_ALIGNMENT.LEFT 

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    if indent_cm != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # --- 3. PROCESSAMENTO DOS DADOS ---
    data_idx = 0
    for i, row_data in enumerate(dados):
        tipo = row_data[0] 
        vals = [str(x) for x in row_data[1:]]
        while len(vals) < NUM_COLUNAS: vals.append("")
        vals = vals[:NUM_COLUNAS]
        
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(ALTURA_LINHA))
        trH.set(qn('w:hRule'), 'atLeast') 
        trPr.append(trH)

        if tipo in ["HEADER", "SUB_HEADER"]:
            trPr.append(OxmlElement('w:tblHeader'))
        else:
            trPr.append(OxmlElement('w:cantSplit'))

        # Header Mesclado
        if tipo == "HEADER":
            c = row.cells[0].merge(row.cells[NUM_COLUNAS-1])
            c.text = vals[0].upper()
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            tcPr = c._element.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), COR_HEADER_BG)
            tcPr.append(shading)
            set_cell_vertical_alignment(c, 'center')
            continue 

        # Colunas Normais (USA lista_larguras AGORA)
        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(lista_larguras[j])) # <--- AQUI A MUDANÇA
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

            cell.text = vals[j]
            p = cell.paragraphs[0]
            
            if j == 0: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = FONTE_TAM
            run.font.color.rgb = RGBColor(0, 0, 0)

            shading = OxmlElement('w:shd')
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')

            if tipo == "SUB_HEADER":
                run.bold = True
                shading.set(qn('w:fill'), COR_SUBHEADER_BG)
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')
            else: 
                run.bold = False
                bg_color = COR_DADOS_BG_PAR if data_idx % 2 == 0 else COR_DADOS_BG_IMPAR
                shading.set(qn('w:fill'), bg_color)
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), 'D9D9D9')
            
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
            tcPr.append(shading)
            set_cell_vertical_alignment(cell, 'center')
            p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        
        if tipo == "DATA_ROW": data_idx += 1

    # --- 4. LEGENDA INFERIOR ---
    if titulo_custom or fonte:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        
        texto_completo = ""
        if titulo_custom:
            texto_completo += str(titulo_custom).strip().rstrip('.') + ". "
        
        txt_fonte = fonte if fonte else "Fonte: CNJ."
        texto_completo += txt_fonte
        run = p.add_run(texto_completo)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)


def adicionar_tabela_6col_simples(document, dados, titulo_custom=None, indent_cm=0, fonte=None):
    """
    Tabela de 6 colunas (Estilo Clean - iGovTIC).
    - Colunas: Tribunal, Nota 24, Nível 24, Nota 25, Nível 25, Variação.
    """
    if not dados: return

    # --- 1. CONFIGURAÇÕES ---
    NUM_COLUNAS = 6
    
    # Larguras Otimizadas (Total ~17cm / 9650 twips)
    # Col 1 (Tribunal): ~2.6cm
    # Cols 2, 4 (Notas): ~1.9cm
    # Cols 3, 5 (Níveis): ~3.2cm (Textos longos como "Satisfatório")
    # Col 6 (Variação): ~2.3cm
    
    larguras = [1500, 1100, 1800, 1100, 1800, 1300]
    # Soma: 8600 twips (se precisar alargar, aumente proporcionalmente)
    # Vamos aumentar um pouco para preencher a página A4 (aprox 9600 útil)
    larguras = [1600, 1300, 2000, 1300, 2000, 1400] # Total 9600
    
    LARGURA_TOTAL_REAL = sum(larguras)
    
    ALTURA_LINHA = 340 
    FONTE_NOME = 'Calibri'
    FONTE_TAM = Pt(10) # Fonte 10 para 6 colunas caberem bem
    ESPACAMENTO_LINHA = 1.15
    
    COR_HEADER_BG = '44546A'    
    COR_SUBHEADER_BG = 'D9D9D9' 
    COR_DADOS_BG_PAR = 'F2F2F2' 
    COR_DADOS_BG_IMPAR = 'D9D9D9' 

    # --- 2. ESTRUTURA XML ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.autofit = False 
    table.alignment = WD_TABLE_ALIGNMENT.LEFT 

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    if indent_cm != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # --- 3. DADOS ---
    data_idx = 0
    for i, row_data in enumerate(dados):
        tipo = row_data[0] 
        vals = [str(x) for x in row_data[1:]]
        while len(vals) < NUM_COLUNAS: vals.append("")
        vals = vals[:NUM_COLUNAS]
        
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(ALTURA_LINHA))
        trH.set(qn('w:hRule'), 'atLeast') 
        trPr.append(trH)

        if tipo in ["HEADER", "SUB_HEADER"]:
            trPr.append(OxmlElement('w:tblHeader'))
        else:
            trPr.append(OxmlElement('w:cantSplit'))

        # HEADER MESCLADO
        if tipo == "HEADER":
            c = row.cells[0].merge(row.cells[NUM_COLUNAS-1])
            c.text = vals[0].upper()
            
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            tcPr = c._element.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), COR_HEADER_BG)
            tcPr.append(shading)
            set_cell_vertical_alignment(c, 'center')
            continue 

        # DADOS
        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(larguras[j]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

            cell.text = vals[j]
            p = cell.paragraphs[0]
            
            # Alinhamento: 1ª Esq, Resto Centro
            if j == 0: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = FONTE_TAM
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Estilização
            shading = OxmlElement('w:shd')
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')

            if tipo == "SUB_HEADER":
                run.bold = True
                shading.set(qn('w:fill'), COR_SUBHEADER_BG)
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')
            else: 
                run.bold = False
                bg_color = COR_DADOS_BG_PAR if data_idx % 2 == 0 else COR_DADOS_BG_IMPAR
                shading.set(qn('w:fill'), bg_color)
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), 'D9D9D9')
            
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
            tcPr.append(shading)
            
            set_cell_vertical_alignment(cell, 'center')
            p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        
        if tipo == "DATA_ROW": data_idx += 1

    # --- 4. LEGENDA INFERIOR ---
    if titulo_custom or fonte:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        
        texto_completo = ""
        if titulo_custom:
            texto_completo += str(titulo_custom).strip().rstrip('.') + ". "
        
        txt_fonte = fonte if fonte else "Fonte: CNJ."
        texto_completo += txt_fonte
        
        run = p.add_run(texto_completo)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)


def adicionar_tabela_comparativo_temas(document, dados, titulo_custom=None, indent_cm=0, fonte=None):
    """
    Tabela de 8 colunas com Seções Intermediárias.
    - Col 1: Tema (Larga)
    - Col 2: Estadual (Média - Destaque)
    - Cols 3-8: TJs (Estreitas)
    """
    if not dados: return

    # --- 1. CONFIGURAÇÕES ---
    NUM_COLUNAS = 8
    
    # DEFINIÇÃO DAS LARGURAS (Ajuste Fino)
    # Total alvo: ~9600 twips (Limite seguro A4)
    
    LARGURA_COL_TEMA = 2500     # Coluna 0: Títulos dos temas
    LARGURA_COL_ESTADUAL = 1400 # Coluna 1: "Estadual" (Mais larga que os TJs)
    LARGURA_RESTO = 950         # Colunas 2-7: Outros TJs (Mais estreitas)
    
    # Monta a lista: [Tema, Estadual, TJ, TJ, TJ, TJ, TJ, TJ]
    larguras = [LARGURA_COL_TEMA, LARGURA_COL_ESTADUAL] + [LARGURA_RESTO] * (NUM_COLUNAS - 2)
    
    LARGURA_TOTAL_REAL = sum(larguras)
    
    ALTURA_LINHA = 300 
    FONTE_NOME = 'Calibri'
    FONTE_TAM = Pt(10)
    ESPACAMENTO_LINHA = 1.05 
    
    # Cores
    COR_HEADER_BG = '44546A'    
    COR_SUBHEADER_BG = 'D9D9D9' 
    COR_SECTION_BG = 'D0CECE'   
    COR_DADOS_BG_PAR = 'F2F2F2' 
    COR_DADOS_BG_IMPAR = 'D9D9D9' 

    # --- 2. ESTRUTURA XML ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.autofit = False 
    table.alignment = WD_TABLE_ALIGNMENT.LEFT 

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    if indent_cm != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # --- 3. PROCESSAMENTO ---
    data_idx = 0
    for i, row_data in enumerate(dados):
        tipo = row_data[0] 
        vals = [str(x) for x in row_data[1:]]
        
        while len(vals) < NUM_COLUNAS: vals.append("")
        vals = vals[:NUM_COLUNAS]
        
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(ALTURA_LINHA))
        trH.set(qn('w:hRule'), 'atLeast') 
        trPr.append(trH)

        if tipo in ["HEADER", "SUB_HEADER"]:
            trPr.append(OxmlElement('w:tblHeader'))
        else:
            trPr.append(OxmlElement('w:cantSplit'))

        # --- A) HEADER MESCLADO ---
        if tipo == "HEADER":
            c = row.cells[0].merge(row.cells[NUM_COLUNAS-1])
            c.text = vals[0].upper()
            
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            tcPr = c._element.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), COR_HEADER_BG)
            tcPr.append(shading)
            set_cell_vertical_alignment(c, 'center')
            continue 

        # --- B) SECTION HEADER ---
        if tipo == "SECTION_HEADER":
            c = row.cells[0].merge(row.cells[NUM_COLUNAS-1])
            c.text = vals[0].upper()
            
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            tcPr = c._element.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), COR_SECTION_BG)
            tcPr.append(shading)
            
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
            
            set_cell_vertical_alignment(c, 'center')
            continue

        # --- C) DADOS E SUB-HEADER ---
        for j, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(larguras[j]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

            cell.text = vals[j]
            p = cell.paragraphs[0]
            
            # Alinhamento: 1ª Esq, Resto Centro
            if j == 0: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.runs[0]
            run.font.name = FONTE_NOME
            run.font.size = FONTE_TAM
            run.font.color.rgb = RGBColor(0, 0, 0)

            shading = OxmlElement('w:shd')
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')

            if tipo == "SUB_HEADER":
                run.bold = True
                shading.set(qn('w:fill'), COR_SUBHEADER_BG)
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')
            else: 
                run.bold = False
                bg_color = COR_DADOS_BG_PAR if data_idx % 2 == 0 else COR_DADOS_BG_IMPAR
                shading.set(qn('w:fill'), bg_color)
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), 'D9D9D9')
            
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
            tcPr.append(shading)
            
            set_cell_vertical_alignment(cell, 'center')
            p.paragraph_format.line_spacing = ESPACAMENTO_LINHA
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        
        if tipo == "DATA_ROW": data_idx += 1

    # --- 4. LEGENDA INFERIOR ---
    if titulo_custom or fonte:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        
        texto_completo = ""
        if titulo_custom:
            texto_completo += str(titulo_custom).strip().rstrip('.') + ". "
        
        txt_fonte = fonte if fonte else "Fonte: CNJ."
        texto_completo += txt_fonte
        
        run = p.add_run(texto_completo)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)


def adicionar_tabela_metas_final(document, dados, titulo_custom=None, indent_cm=0, fonte=None):
    if not dados: return

    # --- 1. DETECÇÃO DINÂMICA DE COLUNAS ---
    # O primeiro elemento de cada linha em 'dados' é o tipo (HEADER, DATA, etc)
    # Portanto, o número de colunas reais da tabela é len(dados[0]) - 1
    NUM_COLUNAS = len(dados[1]) - 1

    # Define as larguras e o índice onde o histórico começa baseado no número de colunas
    if NUM_COLUNAS == 9:
        # Layout para Meta 4 (com coluna 'Grupo/Agregação')
        LARGURAS = [790, 1590, 1530, 1247, 1120, 1120, 1120, 1120, 1120]
        IDX_HISTORICO = 4 # Histórico começa na 5ª coluna
    else:
        # Layout para Metas 1, 2, 3 (padrão 8 colunas)
        LARGURAS = [901, 1338, 1644, 1360, 1360, 1360, 1360, 1360]
        IDX_HISTORICO = 3 # Histórico começa na 4ª coluna

    LARGURA_TOTAL_REAL = sum(LARGURAS)
    ALTURA_LINHA = 392
    
    # Suas cores padronizadas
    COR_CINZA_ESCURO     = '595959' 
    COR_HEADER_HISTORICO = COR_CINZA_ESCURO
    COR_HEADER_MERGED    = COR_CINZA_ESCURO
    COR_HEADER_YEARS     = COR_CINZA_ESCURO
    COR_META_TARGET      = 'D9D9D9' 
    COR_DADOS_PAR        = 'F2F2F2'
    COR_DADOS_IMPAR      = 'FFFFFF'
    COR_DESTAQUE_HEADER  = '91A2B9'
    COR_DESTAQUE_ROW     = 'D5DCE4'

# --- 2. ESTRUTURA DA TABELA ---
    table = document.add_table(rows=0, cols=NUM_COLUNAS)
    table.autofit = False 
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(LARGURA_TOTAL_REAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    if indent_cm != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # --- 3. PROCESSAMENTO DAS LINHAS ---
    indices_inicio_bloco = [] 
    row_header_top = None
    row_header_years = None

    for i, row_data in enumerate(dados):
        tipo = row_data[0]
        vals = [str(x) for x in row_data[1:]]
        while len(vals) < NUM_COLUNAS: vals.append("")
        vals = vals[:NUM_COLUNAS]

        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(ALTURA_LINHA))
        trH.set(qn('w:hRule'), 'atLeast') 
        trPr.append(trH)

        if tipo == "HEADER_TOP":
            row_header_top = row
            for j in range(IDX_HISTORICO):
                estilizar_celula(row.cells[j], vals[j], LARGURAS[j], True, COR_HEADER_MERGED, 'center', True)
            
            c_hist = row.cells[IDX_HISTORICO].merge(row.cells[NUM_COLUNAS - 1]) 
            estilizar_celula(c_hist, vals[IDX_HISTORICO], sum(LARGURAS[IDX_HISTORICO:]), True, COR_HEADER_HISTORICO, 'center', True)
            continue

        if tipo == "HEADER_YEARS":
            row_header_years = row
            for j in range(IDX_HISTORICO):
                estilizar_celula(row.cells[j], "", LARGURAS[j], True, COR_HEADER_MERGED, 'center', True)
            
            for j in range(IDX_HISTORICO, NUM_COLUNAS):
                estilizar_celula(row.cells[j], vals[j], LARGURAS[j], True, COR_HEADER_YEARS, 'center', True)
                # CORREÇÃO 1: Adicionado alinhamento vertical também no cabeçalho
                set_vertical_align(row.cells[j], 'center') 
            
            if row_header_top and row_header_years:
                for col_idx in range(IDX_HISTORICO):
                    txt_orig = row_header_top.cells[col_idx].text
                    merged = row_header_top.cells[col_idx].merge(row_header_years.cells[col_idx])
                    estilizar_celula(merged, txt_orig, LARGURAS[col_idx], True, COR_HEADER_MERGED, 'center', True)
                    set_vertical_align(merged, 'center')
            continue

        if "DATA_ROW" in tipo:
            if tipo == "DATA_ROW_START":
                indices_inicio_bloco.append(len(table.rows) - 1)
            
            idx_real = len(table.rows) - 1
            bg_base = COR_DADOS_PAR if idx_real % 2 == 0 else COR_DADOS_IMPAR
            
            for j, cell in enumerate(row.cells):
                txt, bold, align, remove_bottom, bg = vals[j], False, 'center', False, bg_base
                
                if tipo == "DATA_ROW_START" and j >= IDX_HISTORICO:
                    bg, remove_bottom, bold = COR_META_TARGET, True, True
                
                if j == 1 or j == 2: align = 'center' 

                if tipo == "DATA_ROW_START" and j == NUM_COLUNAS - 1:
                    bg = COR_DESTAQUE_HEADER
                elif j == NUM_COLUNAS - 1:
                    bg = COR_DESTAQUE_ROW
                    bold = True
                
                estilizar_celula(cell, txt, LARGURAS[j], bold, bg, align, False, remove_bottom=remove_bottom)
                # Garante centralização em todas as células de dados
                set_vertical_align(cell, 'center')

# --- 4. MESCLAGEM VERTICAL ---
    # As linhas de dados começam no índice 2 (pois 0=HEADER_TOP e 1=HEADER_YEARS)
    start_row = 2
    end_row = len(table.rows) - 1

    if end_row >= start_row:
        # A. MESCLAGEM TOTAL (Colunas 0 e 1: Meta e Descrição)
        # Elas são mescladas da primeira até a última linha, ignorando grupos internos
        for col_idx in [0, 1]:
            # Pega o texto da primeira célula de dados para preservar
            txt = table.rows[start_row].cells[col_idx].text
            
            # Mescla da linha 2 até a última linha da tabela
            merged_cell = table.rows[start_row].cells[col_idx].merge(table.rows[end_row].cells[col_idx])
            
            # Configura Estilos
            align_h = 'center' if col_idx == 0 else 'left' # Meta centralizada, Desc esquerda
            bold = True if col_idx == 0 else False
            
            estilizar_celula(merged_cell, txt, LARGURAS[col_idx], bold, 'FFFFFF', align_h, False)
            set_vertical_align(merged_cell, 'center')

        # B. MESCLAGEM INTELIGENTE POR GRUPO (Coluna 2 - Apenas se for 9 colunas)
        if NUM_COLUNAS == 9:
            curr_gp = None
            gp_start = start_row
            
            for r in range(start_row, end_row + 1):
                txt_gp = table.rows[r].cells[2].text
                
                # Se o texto do grupo mudou (ex: de "Crimes..." para "Improbidade")
                if txt_gp != curr_gp:
                    # Se não for a primeira passagem, fecha e mescla o grupo anterior
                    if curr_gp is not None:
                        c_g = table.rows[gp_start].cells[2].merge(table.rows[r-1].cells[2])
                        estilizar_celula(c_g, curr_gp, LARGURAS[2], True, 'FFFFFF', 'center', False)
                        set_vertical_align(c_g, 'center')
                    
                    # Inicia o rastreamento do novo grupo
                    curr_gp = txt_gp
                    gp_start = r
            
            # Mescla o último grupo que sobrou no loop (até o final da tabela)
            if gp_start <= end_row:
                c_g = table.rows[gp_start].cells[2].merge(table.rows[end_row].cells[2])
                estilizar_celula(c_g, curr_gp, LARGURAS[2], True, 'FFFFFF', 'center', False)
                set_vertical_align(c_g, 'center')
                
    # --- 5. LEGENDA COM ESPAÇAMENTO (CORREÇÃO 2) ---
    # Removemos o space_after da tabela e aplicamos na legenda (Space Before)
    # Isso evita que o texto da última célula seja empurrado para cima
    
    p = document.add_paragraph() # Cria o parágrafo da legenda
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Aplica o espaçamento de 20pt AQUI, simulando o "space after" da tabela
    p.paragraph_format.space_before = Pt(6) # Espaço visual pequeno entre tabela e legenda
    
    # Adicionamos um espaçamento extra APÓS a legenda para separar do próximo conteúdo
    p.paragraph_format.space_after = Pt(20) 

    texto = f"{str(titulo_custom).strip().rstrip('.')}. " if titulo_custom else ""
    texto += (fonte if fonte else "Fonte: CNJ.")
    run = p.add_run(texto)
    run.font.name, run.font.size, run.font.color.rgb = 'Calibri', Pt(9), RGBColor(0, 0, 0)


def adicionar_tabela_meta_unica_anos(document, nome_meta, anos, valores_meta, valores_resultado, 
                                      titulo_custom=None, indent_cm=0, fonte=None):
    """
    Adiciona uma tabela de meta única com colunas para cada ano com formatação customizada.
    
    Formato:
    ┌──────────────────┬────────┬────────┬────────┬────────┐
    │ Ano              │ 2022   │ 2023   │ 2024   │ 2025   │
    ├──────────────────┼────────┼────────┼────────┼────────┤
    │ Valor da Meta    │   —    │  70%   │  70%   │  70%   │
    ├──────────────────┼────────┼────────┼────────┼────────┤
    │ Resultado        │  60%   │ 64,6%  │  64%   │  65%   │
    └──────────────────┴────────┴────────┴────────┴────────┘
    
    Cores customizadas por linha e coluna:
    - Linha 1 (Anos): BG RGB(89,89,89), Fonte branca, última coluna RGB(68,84,106)
    - Linha 2 (Meta): BG RGB(166,166,166), Col1 fonte branca, resto preto, última col RGB(208,206,206)
    - Linha 3 (Resultado): BG RGB(231,231,231), Fonte preta
    
    Args:
        document: Documento Word
        nome_meta: Nome da meta (ex: "TJMG 5")
        anos: Lista de anos as strings (ex: ['2022', '2023', '2024', '2025'])
        valores_meta: Lista de valores da meta para cada ano (ex: ['—', '70%', '70%', '70%'])
        valores_resultado: Lista de resultados para cada ano (ex: ['60%', '64.6%', '64%', '65%'])
        titulo_custom: Título customizado (opcional)
        indent_cm: Recuo da tabela em cm
        fonte: Fonte a exibir na legenda
    """
    if not anos or not valores_meta or not valores_resultado:
        return
    
    # Cores RGB convertidas para Hex
    # Linha 1: Header (Anos)
    COR_HEADER_LINHA1 = '595959'         # RGB(89,89,89) - Cinza escuro
    COR_FONT_LINHA1 = 'FFFFFF'           # RGB(255,255,255) - Branco
    COR_ULTIMA_COL_LINHA1 = '445468'     # RGB(68,84,106) - Azul
    
    # Linha 2: Valor da Meta
    COR_HEADER_LINHA2 = 'A6A6A6'         # RGB(166,166,166) - Cinza médio
    COR_COL1_FONT_L2 = 'FFFFFF'          # RGB(255,255,255) - Branco (col 1)
    COR_RESTO_FONT_L2 = '000000'         # RGB(0,0,0) - Preto (col 2+)
    COR_ULTIMA_COL_LINHA2 = 'D0CECE'     # RGB(208,206,206) - Cinza claro
    
    # Linha 3: Resultado
    COR_LINHA3 = 'E7E7E7'                # RGB(231,231,231) - Cinza muito claro
    COR_FONT_LINHA3 = '000000'           # RGB(0,0,0) - Preto
    
    # Número de colunas: 1 (rótulo) + N (anos)
    NUM_COLUNAS = len(anos) + 1
    num_anos = len(anos)
    
    # Larguras: todas colunas com 3cm (1701 twips = 3cm * 567 twips/cm)
    LARGURA_COLUNA = 1701  # 3cm
    LARGURAS = [LARGURA_COLUNA] * NUM_COLUNAS
    
    ALTURA_LINHA = 340  # 0.6cm (0.6 * 567 ≈ 340 twips)
    
    # --- Cria a tabela ---
    table = document.add_table(rows=3, cols=NUM_COLUNAS)  # 3 linhas: anos, meta, resultado
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Configura propriedades da tabela
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Define largura fixa
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(LARGURAS)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Aplica recuo se necessário
    if indent_cm != 0:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Cm(indent_cm).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
    
    # Insere tblPr no elemento da tabela se não existir
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # --- LINHA 1: ANOS (Header) ---
    row_anos = table.rows[0]
    set_row_height_at_least(row_anos, ALTURA_LINHA)
    
    # Coluna 1: "Ano" (BG cinza, fonte branca)
    cell = row_anos.cells[0]
    _estilizar_cell_tabela_meta_v2(
        cell, "Ano", LARGURAS[0], True, COR_HEADER_LINHA1, 'left', cor_fonte=COR_FONT_LINHA1
    )
    
    # Colunas 2 a N-1: Anos (BG cinza escuro, fonte branca)
    for idx in range(1, num_anos):
        cell = row_anos.cells[idx]
        _estilizar_cell_tabela_meta_v2(
            cell, anos[idx-1], LARGURAS[idx], True, COR_HEADER_LINHA1, 'center', cor_fonte=COR_FONT_LINHA1
        )
    
    # Última coluna: Último ano (BG azul, fonte branca)
    cell = row_anos.cells[num_anos]
    _estilizar_cell_tabela_meta_v2(
        cell, anos[-1], LARGURAS[num_anos], True, COR_ULTIMA_COL_LINHA1, 'center', cor_fonte=COR_FONT_LINHA1
    )
    
    # --- LINHA 2: VALOR DA META ---
    row_meta = table.rows[1]
    set_row_height_at_least(row_meta, ALTURA_LINHA)
    
    # Coluna 1: "Valor da Meta" (BG cinza, fonte branca)
    cell = row_meta.cells[0]
    _estilizar_cell_tabela_meta_v2(
        cell, "Valor da Meta", LARGURAS[0], False, COR_HEADER_LINHA2, 'left', cor_fonte=COR_COL1_FONT_L2
    )
    
    # Colunas 2 a N-1: Valores meta (BG cinza, fonte preta)
    for idx in range(1, num_anos):
        cell = row_meta.cells[idx]
        _estilizar_cell_tabela_meta_v2(
            cell, valores_meta[idx-1], LARGURAS[idx], False, COR_HEADER_LINHA2, 'center', cor_fonte=COR_RESTO_FONT_L2
        )
    
    # Última coluna: Último valor meta (BG cinza claro, fonte preta)
    cell = row_meta.cells[num_anos]
    _estilizar_cell_tabela_meta_v2(
        cell, valores_meta[-1], LARGURAS[num_anos], False, COR_ULTIMA_COL_LINHA2, 'center', cor_fonte=COR_RESTO_FONT_L2
    )
    
    # --- LINHA 3: RESULTADO ---
    row_resultado = table.rows[2]
    set_row_height_at_least(row_resultado, ALTURA_LINHA)
    
    # Coluna 1: "Resultado" (BG cinza claro, fonte preta)
    cell = row_resultado.cells[0]
    _estilizar_cell_tabela_meta_v2(
        cell, "Resultado", LARGURAS[0], False, COR_LINHA3, 'left', cor_fonte=COR_FONT_LINHA3
    )
    
    # Colunas 2 a N: Valores resultado (BG cinza claro, fonte preta)
    for idx in range(1, NUM_COLUNAS):
        cell = row_resultado.cells[idx]
        _estilizar_cell_tabela_meta_v2(
            cell, valores_resultado[idx-1], LARGURAS[idx], False, COR_LINHA3, 'center', cor_fonte=COR_FONT_LINHA3
        )
    
    # --- Legenda ---
    legenda = document.add_paragraph()
    legenda.paragraph_format.left_indent = Pt(0)
    legenda.paragraph_format.space_before = Pt(6)
    legenda.paragraph_format.space_after = Pt(20)
    legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    titulo_legenda = f"Tabela - {nome_meta}"
    if titulo_custom:
        titulo_legenda = titulo_custom
    
    run = legenda.add_run(titulo_legenda + ". ")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    fonte_texto = fonte if fonte else "Fonte: Dados TJMG"
    run_fonte = legenda.add_run(fonte_texto)
    run_fonte.font.name = 'Calibri'
    run_fonte.font.size = Pt(9)
    run_fonte.font.color.rgb = RGBColor(0, 0, 0)


def _estilizar_cell_tabela_meta_v2(cell, texto, largura_dxa, bold, bg_color, align, cor_fonte='000000'):
    """
    Auxiliar para estilizar célula da tabela meta (versão 2 com cor de fonte customizável).
    
    Args:
        cell: Célula da tabela
        texto: Conteúdo da célula
        largura_dxa: Largura em twips (dxa)
        bold: Se o texto deve estar em bold
        bg_color: Cor de fundo (hex sem #, ex: '595959')
        align: Alinhamento ('left', 'center', 'right')
        cor_fonte: Cor de fonte em hex (ex: 'FFFFFF' para branco, '000000' para preto)
    """
    # Limpa célula
    cell.text = ''
    
    # Configuração de propriedades
    tcPr = cell._element.get_or_add_tcPr()
    
    # Largura
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(largura_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    # Cor de fundo
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    tcPr.append(shading)
    
    # Bordas brancas com 1/2pt
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/2pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')  # Branco
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    # Alinhamento vertical
    tcValign = OxmlElement('w:vAlign')
    tcValign.set(qn('w:val'), 'center')
    tcPr.append(tcValign)
    
    # Conteúdo
    p = cell.paragraphs[0]
    
    # Alinhamento horizontal
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Espaçamento do parágrafo
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(4)
    p.paragraph_format.right_indent = Pt(4)
    
    # Texto com cor customizável
    run = p.add_run(texto)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = bold
    
    # Converte cor hex para RGB
    r = int(cor_fonte[0:2], 16)
    g = int(cor_fonte[2:4], 16)
    b = int(cor_fonte[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


def _estilizar_cell_tabela_meta(cell, texto, largura_dxa, bold, bg_color, align):
    # Limpa célula
    cell.text = ''
    
    # Configuração de propriedades
    tcPr = cell._element.get_or_add_tcPr()
    
    # Largura
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(largura_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    # Cor de fundo
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    tcPr.append(shading)
    
    # Bordas pretas
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Tamanho padrão
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Preto
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    # Alinhamento vertical
    tcValign = OxmlElement('w:vAlign')
    tcValign.set(qn('w:val'), 'center')
    tcPr.append(tcValign)
    
    # Conteúdo
    p = cell.paragraphs[0]
    
    # Alinhamento horizontal
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Espaçamento do parágrafo
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(4)
    p.paragraph_format.right_indent = Pt(4)
    
    # Texto
    run = p.add_run(texto)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


# =============================================================================
# FUNÇÃO PARA ADICIONAR TODAS AS METAS INSTITUCIONAIS
# =============================================================================

def adicionar_todas_metas_institucionais(doc, loader_jn=None):
    """
    Adiciona todas as 55 metas institucionais do TJMG ao documento.
    Lê dados do arquivo: exports/metas_institucionais_2025.xlsx
    
    Cada meta terá:
    - Subtítulo (Heading 3) com ID da meta
    - Descrição (obtida da coluna Nº_Meta)
    - Tabela com histórico (2022, 2023, 2024, 2025)
    - Legenda/Fonte
    """
    import pandas as pd
    import os
    from src.content import static_data
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    def format_cell_meta(cell, bg_color, font_color, border_color='FFFFFF', is_last_col=False, last_col_bg=None, 
                         bold=False, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
        """Formata célula com cor de fundo, fonte e bordas específicas"""
        # Cor de fundo
        tcPr = cell._element.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        
        if is_last_col and last_col_bg:
            shading.set(qn('w:fill'), last_col_bg)
        else:
            shading.set(qn('w:fill'), bg_color)
        tcPr.append(shading)
        
        # Bordas (1/2pt = 2 "eighths of a point")
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 1/2pt = 4 eighths of a point
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        # Cor da fonte e formatação
        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment
            for run in paragraph.runs:
                run.font.color.rgb = font_color
                run.font.bold = bold
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'
    
    print("⚡ Adicionando todas as Metas Institucionais do TJMG (55 metas)...")
    
    # Caminho do arquivo de metas institucionais
    arquivo_metas = 'exports/metas_institucionais_2025.xlsx'
    
    # Verifica se arquivo existe
    if not os.path.exists(arquivo_metas):
        print(f"❌ ERRO: Arquivo {arquivo_metas} não encontrado!")
        return
    
    try:
        # Lê as duas abas do arquivo
        df_valores = pd.read_excel(arquivo_metas, sheet_name='Valores Apurados')
        df_textos = pd.read_excel(arquivo_metas, sheet_name='Textos Metas')
        
        fonte_padrao = "Fonte: Metas Institucionais do TJMG 2025. Dados até 31/12/2025."
        
        # Cores definidas
        header_bg = '595959'  # RGB(89,89,89)
        header_font = RGBColor(255, 255, 255)  # Branco
        last_col_header = '445A6A'  # RGB(68,84,106)
        
        data_row_bg = 'A6A6A6'  # RGB(166,166,166)
        data_row_font_1col = RGBColor(255, 255, 255)  # Branco
        data_row_font_rest = RGBColor(0, 0, 0)  # Preto
        last_col_data = 'D0CECE'  # RGB(208,206,206)
        
        other_rows_bg = 'E7E7E7'  # RGB(231,231,231)
        other_rows_font = RGBColor(0, 0, 0)  # Preto
        
        border_color = 'FFFFFF'  # Branco para bordas
        
        # Itera pelas metas (coluna 'Meta' em df_valores)
        for idx, meta_id in enumerate(df_valores['Meta'].values):
            meta_id = str(meta_id).strip()
            
            # Busca informações da meta em df_textos
            info_texto = df_textos[df_textos['Meta'] == meta_id]
            
            if info_texto.empty:
                print(f"  ⚠ Meta {meta_id} não encontrada em Textos Metas, pulando...")
                continue
            
            descricao = info_texto.iloc[0].get('Nº_Meta', meta_id)
            valor_meta = info_texto.iloc[0].get('Valor da Meta', '')
            
            # --- SUBTÍTULO (Heading 3) ---
            p_titulo = doc.add_paragraph(meta_id, style='Heading 3')
            p_titulo.paragraph_format.space_before = Pt(12)
            p_titulo.paragraph_format.space_after = Pt(6)
            p_titulo.paragraph_format.keep_with_next = True  # Manter título com próximo parágrafo
            
            # Formata o título: cor RGB(162, 22, 18) e tamanho 12pt
            for run in p_titulo.runs:
                run.font.color.rgb = RGBColor(162, 22, 18)
                run.font.size = Pt(12)
            
            # --- DESCRIÇÃO ---
            texto_desc = str(descricao) if pd.notna(descricao) else meta_id
            p_desc = doc.add_paragraph(texto_desc)
            p_desc.paragraph_format.space_after = Pt(12)
            p_desc.paragraph_format.line_spacing = 1.5  # Espaçamento de linha 1,5
            p_desc.paragraph_format.keep_with_next = True  # Manter descrição com tabela
            p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Estiliza o parágrafo: tamanho 12pt
            for run in p_desc.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
            
            # --- TABELA COM HISTÓRICO ---
            dados_meta = df_valores[df_valores['Meta'] == meta_id].iloc[0]
            
            # Monta dados da tabela com 5 colunas (Ano | 2022 | 2023 | 2024 | 2025)
            # e 3 linhas (Ano/Meta/Resultado, valores de Meta, valores de Resultado)
            ano_label = 'Ano'
            meta_label = 'Meta'
            resultado_label = 'Resultado'
            
            # Converte valor_meta para string com percentual
            valor_meta_str = f"{valor_meta:.0f}%" if pd.notna(valor_meta) else ''
            
            # Prepara valores dos anos
            valor_2022 = float(dados_meta.get(2022, 0))
            valor_2023 = float(dados_meta.get(2023, 0))
            valor_2024 = float(dados_meta.get(2024, 0))
            valor_2025 = float(dados_meta.get(2025, 0))
            
            # Monta dados da tabela
            dados_tabela = [
                ('HEADER', ano_label, '2022', '2023', '2024', '2025'),  # Linha 1
                ('DATA', meta_label, valor_meta_str, valor_meta_str, valor_meta_str, valor_meta_str),  # Linha 2 - Meta
                ('DATA', resultado_label, f'{valor_2022:.1f}', f'{valor_2023:.1f}', f'{valor_2024:.1f}', f'{valor_2025:.1f}'),  # Linha 3 - Resultado
            ]
            
            # Cria tabela com 5 colunas
            if len(dados_tabela) >= 3:
                table = doc.add_table(rows=len(dados_tabela), cols=5)
                table.style = 'Table Grid'
                table.autofit = False
                table.allow_autofit = False
                
                # Define largura das colunas via tblGrid (mais confiável)
                tbl = table._element
                
                # Adiciona propriedade para evitar quebra de tabela entre páginas
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                # Não quebra a tabela entre páginas (habilita para cada linha)
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    # Evita quebra de linha da tabela
                    cantSplit = OxmlElement('w:cantSplit')
                    # Remover se já existir
                    for existing in trPr.findall(qn('w:cantSplit')):
                        trPr.remove(existing)
                    trPr.append(cantSplit)
                tblGrid = tbl.find(qn('w:tblGrid'))
                if tblGrid is None:
                    tblGrid = OxmlElement('w:tblGrid')
                    tbl.insert(1, tblGrid)
                else:
                    # Remove gridCol existentes
                    for gridCol in tblGrid.findall(qn('w:gridCol')):
                        tblGrid.remove(gridCol)
                
                # Adiciona 5 gridCol de 1701 twips cada (3cm)
                for i in range(5):
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), '1701')
                    tblGrid.append(gridCol)
                
                # Também define a largura dentro de cada célula
                for row_idx, row in enumerate(table.rows):
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), '1701')
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.insert(0, tcW)
                    
                    # Define altura das linhas em twips (precisamente)
                    # Cabeçalho: 555.59042 twips (arredondado para 556)
                    # Demais linhas: 340.1574 twips (arredondado para 340)
                    if row_idx == 0:
                        set_row_height_flexible(row, 556)  # 555.59042 twips (cabeçalho)
                    else:
                        set_row_height_flexible(row, 340)  # 340.1574 twips (demais linhas)
                
                # Popula e formata tabela
                for row_idx, row_data in enumerate(dados_tabela):
                    row_type = row_data[0]
                    
                    for col_idx in range(5):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell_text = row_data[col_idx + 1] if col_idx + 1 < len(row_data) else ''
                        
                        # Limpa célula e adiciona texto
                        cell.text = ''
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.text = str(cell_text)
                        
                        is_last_col = (col_idx == 4)
                        
                        # Aplica formatação específica por linha
                        if row_type == 'HEADER':  # Linha 1 (Cabeçalho)
                            # Cabeçalho: negrito, tamanho 11, fonte branca, alinhado ao centro
                            if is_last_col:
                                format_cell_meta(cell, header_bg, header_font, border_color, 
                                               is_last_col=True, last_col_bg=last_col_header,
                                               bold=True, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                            else:
                                format_cell_meta(cell, header_bg, header_font, border_color,
                                               bold=True, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                            
                            # Centralizar verticalmente o cabeçalho
                            set_cell_vertical_alignment(cell, 'center')
                        
                        elif row_type == 'DATA':  # Linhas 2 e 3 (Dados)
                            if row_idx == 1:  # Primeira linha de dados (Meta)
                                # Primeira coluna: branca, alinhada à esquerda
                                if col_idx == 0:
                                    format_cell_meta(cell, data_row_bg, data_row_font_1col, border_color,
                                                   bold=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                                # Última coluna com cor especial
                                elif is_last_col:
                                    format_cell_meta(cell, last_col_data, data_row_font_rest, border_color,
                                                   is_last_col=True, last_col_bg=last_col_data,
                                                   bold=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                                # Resto das colunas (preta)
                                else:
                                    format_cell_meta(cell, data_row_bg, data_row_font_rest, border_color,
                                                   bold=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                            
                            else:  # Segunda linha de dados (Resultado)
                                # Cor de fundo RGB(231,231,231), texto preto
                                if col_idx == 0:
                                    format_cell_meta(cell, other_rows_bg, other_rows_font, border_color,
                                                   bold=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                                else:
                                    format_cell_meta(cell, other_rows_bg, other_rows_font, border_color,
                                                   bold=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # --- LEGENDA/FONTE ---
            p_fonte = doc.add_paragraph(fonte_padrao, style='Heading 4')
            p_fonte.paragraph_format.space_before = Pt(6)
            p_fonte.paragraph_format.space_after = Pt(12)
            p_fonte.paragraph_format.keep_with_next = False  # Não manter com próximo para permitir quebra após
            for run in p_fonte.runs:
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # --- ESPAÇAMENTO ENTRE METAS ---
            p_espaco = doc.add_paragraph()
            p_espaco.paragraph_format.space_after = Pt(18)
        
        print(f"✓ {len(df_valores)} metas institucionais adicionadas com sucesso!")
        
    except Exception as e:
        print(f"❌ ERRO ao processar metas institucionais: {str(e)}")
        import traceback
        traceback.print_exc()


