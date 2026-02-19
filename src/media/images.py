from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from src.config import DIR_CANVAS_IMAGES, LARGURA_MAXIMA_IMAGEM

def adicionar_imagem(document, nome_arquivo, titulo="", fonte="Própria", largura_custom=None, recuo_esq=0, space_after=None):    
    """
    Insere imagem com recuo opcional e legenda unificada na parte inferior.
    """
    caminho_completo = DIR_CANVAS_IMAGES / nome_arquivo
    
    # 1. Verificação de existência
    if not caminho_completo.exists():
        p_err = document.add_paragraph(f"[ERRO: Imagem '{nome_arquivo}' não encontrada]")
        p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_err.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        return

    # Define a largura
    largura_final = largura_custom if largura_custom else LARGURA_MAXIMA_IMAGEM

    # --- PARTE ALTERADA: Inserção com Recuo ---
    p_imagem = document.add_paragraph()
    
    # Se houver recuo, aplicamos no parágrafo
    if recuo_esq != 0:
        p_imagem.paragraph_format.left_indent = Cm(recuo_esq)
        # Se tem recuo, geralmente alinhamos à ESQUERDA para o recuo ser respeitado
        p_imagem.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        # Se não tem recuo, mantém o padrão centralizado
        p_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_imagem.paragraph_format.space_after = Pt(2) 
    
    run_img = p_imagem.add_run()
    run_img.add_picture(str(caminho_completo), width=Cm(largura_final))
    
    # --- 2. Legenda Unificada com recuo padrão do documento ---
    p_legenda = document.add_paragraph()
    
    # Define o alinhamento padrão do documento (Justificado ou à Esquerda)
    # Não aplicamos mais o 'recuo_esq' aqui, para que ele siga a margem normal da página
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Configuração de espaçamento após a legenda
    espaco_abaixo = space_after if space_after is not None else 12
    p_legenda.paragraph_format.space_after = Pt(espaco_abaixo)
    
    # 3. Textos da Legenda
    texto_legenda = f"{titulo}. " if titulo else ""
    if fonte:
        texto_legenda += fonte
        
    run_leg = p_legenda.add_run(texto_legenda)
    run_leg.font.name = 'Calibri'
    run_leg.font.size = Pt(9)
    run_leg.font.color.rgb = RGBColor(0, 0, 0)