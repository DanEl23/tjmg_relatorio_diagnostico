import re
import logging
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

# --- IMPORTS PARA PAGINAÇÃO/XML ---
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- IMPORTS DO PROJETO ---
# Certifique-se de que os arquivos __init__.py existem nas pastas
from src.content import static_data 
from src.media import images
from src.tables import builders
from src.extractors.jn_loader import CarregadorJN

# --- CONFIGURAÇÃO DE LOGS ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- CONFIGURAÇÃO VISUAL ---
COR_VINHO = RGBColor(162, 22, 18)
COR_PRETO = RGBColor(0, 0, 0)


# =============================================================================
# FUNÇÕES DE ESTILO E LAYOUT
# =============================================================================

def configurar_layout_pagina(document):
    """ Configura A4, Margens e Distâncias de Cabeçalho/Rodapé """
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # Margens da Página
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    
    # Distâncias
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.25)


def adicionar_paginacao_rodape(document):
    """ 
    Insere numeração de página no rodapé (Alinhado à Direita).
    """
    section = document.sections[0]
    footer = section.footer
    
    # 1. Configura o parágrafo do NÚMERO DA PÁGINA
    if footer.paragraphs:
        p_num = footer.paragraphs[0]
        p_num.text = "" 
    else:
        p_num = footer.add_paragraph()
        
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.paragraph_format.line_spacing = 1.5
    p_num.paragraph_format.space_before = Pt(6)
    p_num.paragraph_format.space_after = Pt(6)

    # 2. Força o Estilo 'Footer'
    try:
        style = document.styles['Footer']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        p_num.style = style
    except KeyError:
        pass

    # 3. Cria o Run e o Campo PAGE via XML
    run = p_num.add_run()
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

    # 4. Adiciona um parágrafo VAZIO DEPOIS (Visual)
    p_vazio = footer.add_paragraph()
    p_vazio.text = ""
    p_vazio.paragraph_format.line_spacing = 1.5 
    p_vazio.paragraph_format.space_before = Pt(6)
    p_vazio.paragraph_format.space_after = Pt(6)


def configurar_estilos_tjmg(document):
    """ Define estilos Heading 1, 2, 3 com a cor Vinho """
    styles = document.styles
    
    def criar_ou_atualizar_estilo(nome, tamanho, recuo, espaco_antes):
        try:
            if nome in styles:
                s = styles[nome]
            else:
                return 

            s.font.name = 'Calibri'
            s.font.size = Pt(tamanho)
            s.font.bold = True
            s.font.color.rgb = COR_VINHO
            
            pf = s.paragraph_format
            pf.space_before = Pt(espaco_antes)
            pf.space_after = Pt(6)
            pf.left_indent = Cm(recuo)
        except: pass

    # Heading 1: Tamanho 16, Recuo de 1.25 cm
    criar_ou_atualizar_estilo('Heading 1', 18, 1.25, 15)
    
    # Heading 2 e 3: Tamanho 16, Sem recuo
    criar_ou_atualizar_estilo('Heading 2', 16, 0.0, 15)
    criar_ou_atualizar_estilo('Heading 3', 16, 0.0, 15)


def inserir_capa(document, pasta_resources):
    """ Insere a imagem de capa se ela existir """
    caminho_capa = pasta_resources / "capa_relatorio.png"
    
    if caminho_capa.exists():
        print("🖼️ Inserindo Capa...")
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(caminho_capa), width=Cm(21.0))
        document.add_section(WD_SECTION.NEW_PAGE)
    else:
        print(f"⚠️ Capa não encontrada em: {caminho_capa}")


def adicionar_texto_com_negrito(paragrafo, texto, cor_rgb=RGBColor(0,0,0), tamanho=12):
    """
    Processa o texto procurando por trechos entre asteriscos (*texto*).
    """
    partes = re.split(r'(\*[^*]+\*)', texto)
    
    for parte in partes:
        if not parte: continue 
        
        run = paragrafo.add_run()
        
        # Se for um trecho entre asteriscos (*negrito*)
        if parte.startswith('*') and parte.endswith('*') and len(parte) > 2:
            texto_limpo = parte[1:-1]
            run.text = texto_limpo
            run.bold = True
        else:
            run.text = parte
            run.bold = False
            
        run.font.name = 'Calibri'
        run.font.size = Pt(tamanho)
        run.font.color.rgb = cor_rgb


def adicionar_pagina_sumario_visual(doc, doc_orig):
    """ Gera o sumário visual copiando do modelo """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    run = p.add_run("SUMÁRIO")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COR_VINHO
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(24)

    for para in doc_orig.paragraphs:
        txt = para.text.strip()
        if not txt: continue
        match = re.match(r'^(\d+(\.\d+)*)\.?\s+(.*)', txt)
        if match:
            num = match.group(1)
            t = match.group(3).split('...')[0].strip()
            lvl = num.count('.') + 1
            
            p_item = doc.add_paragraph()
            
            if lvl == 1: indent = Cm(0)
            elif lvl == 2: indent = Cm(0.42)
            elif lvl == 3: indent = Cm(0.85)
            else: indent = Cm(0.85)
                
            p_item.paragraph_format.left_indent = indent
            p_item.paragraph_format.line_spacing = 1.5
            p_item.paragraph_format.space_before = Pt(6)
            p_item.paragraph_format.space_after = Pt(5)
            
            run = p_item.add_run(f"{num} {t.upper() if lvl==1 else t}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COR_PRETO
            run.font.name = 'Calibri'


# =============================================================================
# PROCESSAMENTO DE RECURSOS (TABELAS, IMAGENS)
# =============================================================================

def preparar_dados_tabela_metas(nome_meta):
    import pandas as pd
    from src.content import static_data
    
    # 1. Carregamento do Excel
    try:
        df_2025 = pd.read_excel("exports/resultados_cnj.xlsx")
        df_filtrado = df_2025[df_2025['Meta'].astype(str).str.contains(nome_meta, na=False)]
    except:
        df_filtrado = pd.DataFrame()

    info_meta = static_data.HISTORICO_METAS_CNJ.get(nome_meta)
    if not info_meta: return []

    # --- CASO A: META COM SUBGRUPOS (Ex: Meta 4) ---
    if 'grupos' in info_meta:
        # Cabeçalho para 9 colunas
        dados_finais = [
            ["HEADER_TOP", "META", "DESCRIÇÃO", "GRUPO", "INSTÂNCIA", "HISTÓRICO"],
            ["HEADER_YEARS", "", "", "", "", "2021", "2022", "2023", "2024", "2025*"]
        ]
        
        for grupo in info_meta['grupos']:
            nome_visual = grupo['nome']
            chave_excel = grupo.get('chave_busca', nome_visual)
            
            # --- CORREÇÃO AQUI: Lógica para pegar Lista ou Valor Único ---
            if 'objetivos_anos' in grupo:
                lista_objs = grupo['objetivos_anos'] # Usa a lista de 5 anos
            elif 'objetivo' in grupo:
                lista_objs = [grupo['objetivo']] * 5 # Repete o valor único
            else:
                lista_objs = ["---"] * 5 # Preenche com traços se não houver meta
            
            # Adiciona a linha de Meta (Cinza) com os objetivos corretos
            dados_finais.append(["DATA_ROW_START", nome_meta, info_meta['descricao'], nome_visual, "Meta"] + lista_objs)
            
            # Processamento dos resultados (Instâncias)
            for instancia, valores_anos in grupo['dados'].items():
                valor_2025 = "---"
                if not df_filtrado.empty:
                    inst_busca = "Total" if instancia == "Geral" else instancia
                    termo_busca = f"{inst_busca} - {chave_excel}"
                    
                    match = df_filtrado[df_filtrado['Categoria'].astype(str).str.strip() == termo_busca]
                    if not match.empty:
                        valor_2025 = str(match.iloc[0]['Resultado'])
                
                # Monta a linha de dados
                dados_finais.append(["DATA_ROW", nome_meta, info_meta['descricao'], nome_visual, instancia] + valores_anos + [valor_2025])
        
        return dados_finais

    # --- CASO B: META SIMPLES (Ex: Meta 1, 2, 3) ---
    else:
        # Cabeçalho para 8 colunas
        dados_finais = [
            ["HEADER_TOP", "META", "DESCRIÇÃO", "INSTÂNCIA", "HISTÓRICO"],
            ["HEADER_YEARS", "", "", "", "2021", "2022", "2023", "2024", "2025*"]
        ]

        if 'objetivos_anos' in info_meta:
            lista_objs = info_meta['objetivos_anos']
        elif 'objetivo' in info_meta:
            lista_objs = [info_meta['objetivo']] * 5
        else:
            lista_objs = None # Para metas sem objetivo explícito (Meta 3)

        if lista_objs:
            dados_finais.append(["DATA_ROW_START", nome_meta, info_meta['descricao'], ""] + lista_objs)

        for instancia, valores_anos in info_meta.get('dados_passados', {}).items():
            valor_2025 = "---"
            if not df_filtrado.empty:
                # 1. Definição do termo de busca exato
                termo_busca = instancia
                # Pequeno ajuste de compatibilidade: se no static for "Geral", busca "Total" no Excel
                if instancia == "Geral": termo_busca = "Total"

                # 2. Busca Estrita (Removemos o 'OR Total')
                # Isso garante que 1º Grau só pegue dados de 1º Grau.
                match = df_filtrado[
                    df_filtrado['Categoria'].astype(str).str.strip() == termo_busca
                ]
                
                if not match.empty:
                    valor_2025 = str(match.iloc[0]['Resultado'])
            
            dados_finais.append(["DATA_ROW", nome_meta, info_meta['descricao'], instancia] + valores_anos + [valor_2025])
            
        return dados_finais        


def processar_recurso(doc, chave, item, loader_jn=None):
    """ Processa tabelas, imagens e gráficos baseado no mapa de recursos """
    tipo = item["tipo"]
    dados = item.get("dados")
    
    titulo_real = item.get("titulo", chave) 
    fonte_custom = item.get("fonte_custom")

    print(f"⚡ Inserindo Recurso: {titulo_real} (Tipo: {tipo})")

    # === IMAGENS ===
    if tipo == "IMAGEM":
        images.adicionar_imagem(
            doc, item["arquivo"], titulo=titulo_real, 
            fonte=item.get("fonte", "Própria"),
            largura_custom=item.get("largura"),
            recuo_esq=item.get("recuo_esq", 0)
        )
    
    # === TABELAS DE ORÇAMENTO ===
    elif tipo == "TABELA_ORCAMENTO_CONJUNTO":
        builders.adicionar_tabela_orcamento_conjunto(doc, dados)
    elif tipo == "TABELA_ORCAMENTO_CONJUNTO_COMPARACAO":
        builders.adicionar_tabela_orcamento_detalhada(doc, dados)
    elif tipo == "TABELA_ORCAMENTO":
        builders.adicionar_tabela_orcamento(
            doc, titulo_vindo_do_word=titulo_real, dados=dados, 
            numero_tabela=item.get("num", "09"),
            titulo_custom=item.get("titulo")
        )
    elif tipo == "TABELA_SIMPLES_3COL":
            recuo_custom = item.get("recuo_esq", 0)
            fonte_custom = item.get("fonte_custom") # Pega a fonte do static_data

            builders.adicionar_tabela_simples_3col(
                doc, 
                item['dados'], 
                titulo_custom=titulo_real, # Passa o nome da tabela
                indent_cm=recuo_custom,
                fonte=fonte_custom          # Passa a fonte para o builder juntar
            )

    elif tipo == "TABELA_4COL_SIMPLES":
            recuo_custom = item.get("recuo_esq", 0)
            fonte_custom = item.get("fonte_custom")
            larguras_custom = item.get("larguras") # Pega se existir

            builders.adicionar_tabela_4col_simples(
                doc, 
                item['dados'], 
                titulo_custom=titulo_real,
                indent_cm=recuo_custom,
                fonte=fonte_custom,
                larguras=larguras_custom # Passa para o builder
            )

    elif tipo == "TABELA_6COL_SIMPLES":
            recuo_custom = item.get("recuo_esq", 0)
            fonte_custom = item.get("fonte_custom")

            builders.adicionar_tabela_6col_simples(
                doc, 
                item['dados'], 
                titulo_custom=titulo_real,
                indent_cm=recuo_custom,
                fonte=fonte_custom
            )

    elif tipo == "TABELA_COMPARATIVO_TEMAS":
            recuo_custom = item.get("recuo_esq", 0)
            fonte_custom = item.get("fonte_custom")

            builders.adicionar_tabela_comparativo_temas(
                doc, 
                item['dados'], 
                titulo_custom=titulo_real,
                indent_cm=recuo_custom,
                fonte=fonte_custom
            )

    elif tipo == "TABELA_METAS_DINAMICA":
            # O nome_recurso seria "Meta 1"
            dados_processados = preparar_dados_tabela_metas(titulo_real)
            
            if dados_processados:
                builders.adicionar_tabela_metas_final(
                    doc, 
                    dados_processados, 
                    titulo_custom=item.get("titulo_legenda", titulo_real),
                    indent_cm=item.get("recuo_esq", 0),
                    fonte=item.get("fonte_custom")
                )

    # === TABELAS ESPECÍFICAS (Manuais) ===
    elif tipo == "TABELA_PROCESSOS": builders.adicionar_tabela_processos(doc, dados, texto_legenda=titulo_real)
    elif tipo == "TABELA_ATOS": builders.adicionar_tabela_atos(doc, dados)
    elif tipo == "TABELA_AREAS": builders.adicionar_tabela_areas(doc, dados)
    elif tipo == "TABELA_ESTRUTURA": builders.adicionar_tabela_estrutura(doc, dados)
    elif tipo == "TABELA_COMARCAS": builders.adicionar_tabela_comarcas(doc, dados)
    elif tipo == "TABELA_NUCLEOS": builders.adicionar_tabela_nucleos(doc, dados)
    elif tipo == "TABELA_CIDADES": builders.adicionar_tabela_cidades(doc, dados)
        
# === TABELA JUSTIÇA EM NÚMEROS (CONFERIDA E ALINHADA) ===
    elif tipo == "TABELA_JUSTICA_NUMEROS":
        if not loader_jn:
            print("❌ Erro: Loader JN não foi inicializado.")
            return

        # 1. LISTA DE MÉTRICAS (CHAVES INTERNAS DO LOADER)
        # Importante: A ordem aqui deve ser IDÊNTICA à ordem da lista de títulos abaixo.
        metricas = [
            "municipios", #1
            "pop_sede_perc", #2 
            "unidades_jud", #3
            "ranking_tjmg", #4
            "magistrados", #5
            "forca_trabalho", #6 
            "despesa_total", #7
            "despesa_hab", #8
            "custo_magistrado", #9 
            "custo_servidor", #10
            "perc_cargos_vagos_mag", #11
            "perc_serv_adm", #12
            "casos_novos", #13 
            "casos_pendentes", #14 
            "cn_100k_hab", #15
            "ipm", #16
            "ipsjud", #17
            "perc_serv_jud_1grau", #18 
            "iad", #19
            "perc_eletr", #20 
            "perc_unidades_j100", #21
            "nucleos_40", #22 
            "balcao_virtual", #23 
            "cn_mag_1", #24
            "cn_mag_2", #25
            "cn_serv_1", #26
            "cn_serv_2", #27
            "carga_mag_1", #28
            "carga_mag_2", #29
            "carga_serv_1", #30
            "carga_serv_2", #31
            "ipm_1", #32
            "ipm_2", #33
            "ips_1", #34
            "ips_2", #35
            "ind_cn_eletr", #36 
            "perc_eletr_1", #37
            "perc_eletr_2", #38
            "iad_1", #39
            "iad_2", #40
            "tc_total", #41
            "tc_liq", #42
            "tc_1", #43
            "tc_2", #44
            "tc_conhec", #45 
            "tc_exec", #46
            "rin_geral", #47
            "rx_geral", #48
            "rin_1", #49
            "rin_2", #50
            "rx_1", #51
            "rx_2", #52
            "perc_pend_exec_estoque", #53 
            "pend_exec_fiscal", #54
            "tc_exec_fiscal", #55
            "cejusc", #56
            "ic_geral", #57
            "ic_1", #58
            "ic_2", #59
            "tempo_sent_1", #60
            "tempo_sent_2", #61
            "tempo_giro", #62
            "tempo_fisico", #63
            "tempo_eletr", #64
            "cn_crim", #65
            "cp_crim", #66
            "ipc_jus", #67
            "ipc_jus_1", #68
            "ipc_jus_2", #69
            "ipm_meta", #70
            "ips_meta", #71
            "tcl_meta"#72
        ]
        
        # 2. LISTA DE TÍTULOS (NOMES QUE SERÃO IMPRESSOS NO WORD)
        titulos = [
            "Nº de municípios-sede", #1
            "Percentual da população em munícipios-sede", #2
            "Nº de unidades judiciárias (Estrutura de 1º grau)", #3
            "Classificação do TJMG dentro do Grupo ‘Grande Porte’", #4
            "Nº de magistrados", #5
            "Força de trabalho (servidores e auxiliares) (*)", #6
            "Despesa total da justiça (Bilhões)", #7
            "Despesa total por habitante, incluindo custo com inativos (Reais)", #8
            "Custo médio mensal com magistrados (Milhões)", #9
            "Custo médio mensal com servidores (Milhões)", #10
            "Percentual de cargos vagos de magistrados", #11
            "Percentual de servidores lotados na área administrativa", #12
            "Casos novos", #13
            "Casos pendentes", #14
            "Casos novos por 100 mil habitantes", #15
            "Índice de produtividade dos magistrados", #16
            "Índice de produtividade de servidores da área judiciária", #17
            "Percentual de servidores (as) na área judiciária de primeiro grau", #18
            "Índice de atendimento à demanda (Geral)", #19
            "Percentual de casos novos eletrônicos", #20
            "Percentual de unidades judiciárias de primeiro grau com Juízo 100% Digital", #21
            "Quantidade de Núcleos de Justiça 4.0", #22
            "Quantidade de Balcões Virtuais instalados", #23
            "Casos novos por magistrados - 1º grau", #24
            "Casos novos por magistrados - 2º grau", #25
            "Casos novos por servidor da área judiciária – 1º grau", #26
            "Casos novos por servidor da área judiciária – 2º grau", #27
            "Carga de trabalho do magistrado – 1º grau", #28
            "Carga de trabalho do magistrado – 2º grau", #29
            "Carga de trabalho do servidor da área judiciária – 1º grau", #30
            "Carga de trabalho do servidor da área judiciária – 2º grau", #31
            "Índice de produtividade dos magistrados – 1º grau", #32
            "Índice de produtividade dos magistrados – 2º grau", #33
            "Índice de produtividade dos servidores da área judiciária – 1º grau", #34
            "Índice de produtividade dos servidores da área judiciária – 2º grau", #35
            "Índice de casos novos eletrônicos", #36
            "Índice de casos novos eletrônicos – 1º grau", #37
            "Índice de casos novos eletrônicos – 2º grau", #38
            "Índice de atendimento à demanda – 1º grau", #39
            "Índice de atendimento à demanda – 2º grau", #40
            "Taxa de congestionamento Total", #41
            "Taxa de congestionamento líquida", #42
            "Taxa de congestionamento – 1º grau", #43
            "Taxa de congestionamento – 2º grau", #44
            "Taxa de congestionamento na fase de conhecimento", #45
            "Taxa de congestionamento na fase de execução", #46
            "Índice de recorribilidade interna (Geral)", #47
            "Índice de recorribilidade externa (Geral)", #48
            "Recorribilidade interna – 1º grau (Conhecimento)", #49
            "Recorribilidade interna – 2º grau (**)", #50
            "Recorribilidade externa – 1º grau (Conhecimento)", #51
            "Recorribilidade externa – 2º grau (**)", #52
            "Percentual de casos pendentes de execução em relação ao estoque total de processos", #53
            "Total de execuções fiscais pendentes", #54
            "Taxa de congestionamento na execução fiscal", #55
            "Centros judiciários de solução de conflitos na justiça estadual", #56
            "Índice de conciliação", #57
            "Índice de conciliação, 1º grau", #58
            "Índice de conciliação 2º grau", #59
            "Tempo médio até a sentença no 1º grau", #60
            "Tempo médio até a sentença no 2º grau", #61
            "Tempo de giro do acervo", #62
            "Tempo médio dos processos físicos pendentes", #63
            "Tempo médio dos processos eletrônicos pendentes", #64
            "Casos novos criminais, excluídas as execuções penais", #65
            "Casos pendentes criminais, excluídas as execuções penais", #66
            "Resultado do IPC-Jus total por tribunal (incluída a área administrativa)", #67
            "Resultado do IPC-Jus da área judiciária, por instância e tribunal. 1º grau", #68
            "Resultado do IPC-Jus da área judiciária, por instância e tribunal. 2º grau", #69
            "Índice de produtividade dos magistrados (IPM) realizado x necessário para que tribunal atinja IPC-Jus de 100%.", #70
            "Índice de produtividade dos servidores (IPS) realizado x necessário para que tribunal atinja IPC-Jus de 100%.", #71
            "Taxa de congestionamento líquida (TCL) realizado x resultado da consequência se tribunal atingisse IPC-Jus 100%. TCL realizado" #72
        ]
        
        # 3. GERAÇÃO DA TABELA
        dados_tabela = loader_jn.obter_dados_tabela(
            tribunal_sigla='TJMG',
            lista_metricas_amigaveis=metricas,
            anos=[2019, 2020, 2021, 2022, 2023, 2024],
            titulos_linhas=titulos
        )
        
        if dados_tabela:
            # Insere cabeçalho mesclado
            if dados_tabela[0][0] != 'HEADER_MERGE':
                dados_tabela.insert(0, ['HEADER_MERGE', 'RELATÓRIO JUSTIÇA EM NÚMEROS (CNJ) | DADOS DO TJMG'] + ['']*6)
            
            # Aplica recuo customizado do static_data
            recuo_custom = item.get("recuo_esq", -0.5)
            builders.adicionar_tabela_justica_numeros(
                doc, dados_tabela, texto_legenda=titulo_real, indent_cm=recuo_custom
            )
    # === TABELA GENÉRICA ===
    elif tipo == "TABELA_GENERICA":
        builders.adicionar_tabela_generica(doc, titulo_tabela=titulo_real, dados=dados, fonte=fonte_custom)
    
    # === TODAS AS METAS INSTITUCIONAIS ===
    elif tipo == "METAS_INSTITUCIONAIS":
        builders.adicionar_todas_metas_institucionais(doc, loader_jn=loader_jn)
    
    # Espaçamento final
    espaco_final = item.get("space_after", 6) 
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(espaco_final)


# =============================================================================
# FUNÇÃO PRINCIPAL (MAIN)
# =============================================================================

def gerar_relatorio_completo(caminho_base_dummy, output_path, mapa_recursos=None):
    """ Gera o relatório DO ZERO (Blank Document). """
    print(f"--- 🚀 Iniciando Geração (Modo Zero-Base) ---")

    # 1. SETUP DE DIRETÓRIOS (CAMINHOS ABSOLUTOS SEGUROS)
    # Usa a localização deste arquivo (src/core/generator.py) para achar a raiz
    base_dir = Path(__file__).resolve().parent.parent.parent
    
    # Define pastas
    pasta_raw = base_dir / "data" / "raw"
    pasta_processed = base_dir / "data" / "processed"
    pasta_resources = base_dir / "resources"
    
    # Define arquivos
    caminho_conteudo = pasta_processed / "Conteudo_Fonte.docx"
    caminho_sumario = pasta_processed / "Sumario_Modelo.docx"
    arquivo_dados = pasta_raw / "JN_15-Jan-2026.csv"
    arquivo_manual = pasta_raw / "dados_manuais.csv"
    
    
    print(f"📂 Raiz do Projeto: {base_dir}")
    print(f"📂 Procurando CSV em: {arquivo_dados}")

    # Validação do CSV
    if not arquivo_dados.exists():
        print(f"❌ ERRO CRÍTICO: Arquivo de dados não encontrado em {arquivo_dados}")
        print("   -> O relatório será gerado, mas a Tabela Justiça em Números falhará.")
    
    # 2. INICIALIZAÇÃO DE CARREGADORES
    loader_jn = CarregadorJN(arquivo_dados, arquivo_manual)

    # 3. CRIAÇÃO DO DOCUMENTO
    doc_final = Document()
    configurar_layout_pagina(doc_final)
    configurar_estilos_tjmg(doc_final)
    adicionar_paginacao_rodape(doc_final)

    # 4. CAPA
    inserir_capa(doc_final, pasta_resources)

    # 5. SUMÁRIO (Visual)
    try:
        # Fallback de caminho para sumário
        if not caminho_sumario.exists():
             # Tenta achar no diretório pai do dummy path (caso de teste)
             caminho_sumario = Path(caminho_base_dummy).parent / "Sumario_Modelo.docx"

        doc_sumario_orig = Document(caminho_sumario)
        print("📋 Gerando página de Sumário...")
        adicionar_pagina_sumario_visual(doc_final, doc_sumario_orig)
        doc_final.add_page_break()
    except Exception as e:
        print(f"⚠️ Erro ao ler Sumario_Modelo ({e}). Pulando sumário.")

    # 6. PROCESSAMENTO DO CONTEÚDO
    try:
        # Fallback de caminho para conteúdo
        if not caminho_conteudo.exists():
            caminho_conteudo = Path(caminho_base_dummy).parent / "Conteudo_Fonte.docx"
            
        doc_fonte = Document(caminho_conteudo)
    except:
        print("❌ Erro fatal: Conteudo_Fonte.docx não encontrado."); return

    mapa = static_data.MAPA_RECURSOS
    
    # Variáveis de Estado (Listas)
    em_lista_numerica = False
    em_lista_marcadores = False

    print("--- Processando Texto ---")
    for para in doc_fonte.paragraphs:
        texto = para.text.strip()
        if not texto: continue
        
        # Filtros e Lógica de Listas
        if texto.upper() == "SUMÁRIO" or re.match(r'^\d+$', texto): continue
        if "[INICIAR_LISTA_NUMERICA]" in texto: em_lista_numerica = True; continue 
        if "[FINALIZAR_LISTA_NUMERICA]" in texto: em_lista_numerica = False; continue
        if "[INICIAR_LISTA_MARCADORES]" in texto: em_lista_marcadores = True; continue
        if "[FINALIZAR_LISTA_MARCADORES]" in texto: em_lista_marcadores = False; continue
        if "[QUEBRA_PAGINA]" in texto: doc_final.add_page_break(); continue
        if texto.startswith("#"):
            texto_titulo = texto.lstrip("#").strip()
            p_novo = doc_final.add_paragraph()
            p_novo.paragraph_format.space_before = Pt(12)
            run = p_novo.add_run(texto_titulo)
            run.font.name = 'Calibri'  # Nome da fonte desejada
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri') # Garante a aplicação no Word
            run.font.color.rgb = RGBColor(162, 22, 18) 
            run.font.bold = True
            run.font.size = Pt(12)
            continue

        # --- RECURSOS VISUAIS ---
        if texto in mapa:
            # Passamos o loader_jn para a função processar
            processar_recurso(doc_final, texto, mapa[texto], loader_jn=loader_jn)
            continue

        # --- TÍTULOS (Headings) ---
        match = re.match(r'^\s*(\d+(?:\.\d+)*\.?)\s+(.*)', texto)
        eh_titulo_valido = False
        
        if match:
            prefixo = match.group(1).strip()
            titulo_texto = match.group(2).strip()
            
            tem_ponto = '.' in prefixo
            segmentos = prefixo.replace('.', ' ').split()
            tem_numero_grande = any(len(seg) > 2 for seg in segmentos)
            
            if tem_ponto and not tem_numero_grande:
                eh_titulo_valido = True
                num_limpo = prefixo.rstrip('.')
                nivel = num_limpo.count('.') + 1
                if nivel > 3: nivel = 3
                
                if nivel == 1: texto_final_titulo = f"{num_limpo}. {titulo_texto}"
                else: texto_final_titulo = f"{num_limpo} {titulo_texto}"

                print(f"🔖 Título Detectado: {texto_final_titulo}")
                h = doc_final.add_heading(texto_final_titulo, level=nivel)
                if h.runs: 
                    h.runs[0].font.color.rgb = COR_VINHO
                    h.runs[0].font.name = 'Calibri'
        
        if eh_titulo_valido: continue

        # --- TEXTO DESTAQUE ---
        if texto.startswith('#'):
            texto_limpo = texto.lstrip('#').strip()
            p = doc_final.add_paragraph(texto_limpo)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.runs[0] if p.runs else p.add_run(texto_limpo)
            run.font.name = 'Calibri'; run.font.size = Pt(18)
            run.bold = True; run.font.color.rgb = COR_VINHO
            continue

        # --- TEXTO COMUM ---
        p = doc_final.add_paragraph() 
        
        if em_lista_numerica:
            try: p.style = 'List Number'
            except: pass 
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0 
            p.paragraph_format.left_indent = Cm(1.27) 
            p.paragraph_format.first_line_indent = Cm(-0.63)
            
        elif em_lista_marcadores:
            try: p.style = 'List Bullet' # Tenta aplicar o estilo de bolinha padrão
            except: pass
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5 
            if "[MARK_NIVEL_DOIS]" in texto:
                texto = texto.replace("[MARK_NIVEL_DOIS]", "").strip()
                p.paragraph_format.left_indent = Cm(2.54)
            else:
                p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.63)

        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5  
            p.paragraph_format.space_after = Pt(10) 
        
        if "[ICON_CHECK]" in texto:
            # 1. Remove a tag do texto
            texto = texto.replace("[ICON_CHECK]", "").strip()
            
            # 2. Adiciona o ícone visualmente
            run_icon = p.add_run("\u2714  ") # Checkmark Unicode + espaço
            run_icon.font.name = "Segoe UI Symbol"
            run_icon.font.size = Pt(11)
            run_icon.font.color.rgb = RGBColor(0, 150, 0) # Verde
        # =================================================================
        if texto.startswith('[MARK_NOTA]'):  # <--- SEU NOVO BLOCO AQUI
            texto_limpo = texto.replace("[MARK_NOTA]", "").strip()
            p.paragraph_format.line_spacing = 1.0 
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6) 
            
            run = p.add_run(texto_limpo)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'  
            texto = ""     


        # Chama a função que escreve o restante do texto em negrito se precisar
        adicionar_texto_com_negrito(p, texto, cor_rgb=COR_PRETO, tamanho=12)
        
    # 7. SALVAR
    try:
        doc_final.save(output_path)
        print(f"✅ Relatório salvo em: {output_path}")
    except Exception as e:
        print(f"❌ Erro ao salvar: {e}")

if __name__ == "__main__":
    # Teste isolado
    print("Testando gerador isoladamente...")
    try:
        gerar_relatorio_completo("template_dummy.docx", "relatorio_teste.docx")
    except Exception as e:
        print(f"Erro no teste: {e}")