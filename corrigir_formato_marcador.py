#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def corrigir_formato_marcador():
    """Corrige o formato do marcador removendo o prefixo MAPA_RECURSOS::"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("CORRIGINDO FORMATO DO MARCADOR")
    print("=" * 100)
    
    # Procura por "MAPA_RECURSOS::METAS_INSTITUCIONAIS"
    indices_para_corrigir = []
    
    for idx, p in enumerate(doc.paragraphs):
        if "MAPA_RECURSOS::METAS_INSTITUCIONAIS" in p.text:
            indices_para_corrigir.append(idx)
            print(f"\nEncontrado em [{idx}]: {p.text}")
    
    # Corrige para apenas "METAS_INSTITUCIONAIS"
    for idx in indices_para_corrigir:
        p = doc.paragraphs[idx]
        p.text = "METAS_INSTITUCIONAIS"
        print(f"Corrigido para: METAS_INSTITUCIONAIS")
    
    # Salva
    doc.save('data/raw/Conteudo_Fonte.docx')
    
    print(f"\n✓ {len(indices_para_corrigir)} marcador(es) corrigido(s)!")
    print("✓ Documento salvo!")
    print("\nPróximo passo: Execute novamente o gerador com:")
    print("  python main.py --saida Relatorio_Com_Metas_Final_v2.docx")

if __name__ == '__main__':
    corrigir_formato_marcador()
