#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.oxml.ns import qn

def verificar_break_prevention():
    """Verifica se as propriedades de quebra de página foram aplicadas"""
    doc = Document('teste_metas_institucionais.docx')
    
    print("=" * 100)
    print("VERIFICAÇÃO DE PROPRIEDADES DE QUEBRA DE PÁGINA")
    print("=" * 100)
    
    # Verifica primeira meta
    print("\nVerificando primeira meta (TJMG 5):\n")
    
    # Procura pelos parágrafos da primeira meta
    para_idx = 0
    meta_count = 0
    
    for p_idx, p in enumerate(doc.paragraphs[:20]):
        if p.text.startswith('TJMG'):
            meta_count += 1
            if meta_count == 1:
                print(f"Título (parágrafo {p_idx}): {p.text}")
                pPr = p._element.pPr
                keep_next = pPr.find(qn('w:keepNext')) if pPr is not None else None
                print(f"  keep_with_next: {'SIM ✓' if keep_next is not None else 'NÃO'}")
                
                # Próximo parágrafo (descrição)
                p_desc = doc.paragraphs[p_idx + 1]
                print(f"\nDescrição (parágrafo {p_idx + 1}): {p_desc.text[:60]}...")
                pPr_desc = p_desc._element.pPr
                keep_next_desc = pPr_desc.find(qn('w:keepNext')) if pPr_desc is not None else None
                print(f"  keep_with_next: {'SIM ✓' if keep_next_desc is not None else 'NÃO'}")
                
                # Primeira tabela
                tabela = doc.tables[0]
                print(f"\nTabela (primeira tabela do documento):")
                print(f"  Linhas: {len(tabela.rows)}")
                
                # Verifica cada linha
                for row_idx, row in enumerate(tabela.rows):
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    cantSplit = trPr.find(qn('w:cantSplit'))
                    status = "SIM ✓" if cantSplit is not None else "NÃO"
                    print(f"  Linha {row_idx + 1} - cantSplit (não quebra): {status}")
                
                # Legenda
                for p_idx2 in range(p_idx + 2, min(p_idx + 10, len(doc.paragraphs))):
                    p_leg = doc.paragraphs[p_idx2]
                    if p_leg.text.startswith('Fonte:'):
                        print(f"\nLegenda (parágrafo {p_idx2}): {p_leg.text[:60]}...")
                        pPr_leg = p_leg._element.pPr
                        keep_next_leg = pPr_leg.find(qn('w:keepNext')) if pPr_leg is not None else None
                        print(f"  keep_with_next: {'SIM ✓' if keep_next_leg is None else 'NÃO (permitindo quebra após)'}")
                        break
                
                break
    
    print("\n" + "=" * 100)
    print("✓ VERIFICAÇÃO CONCLUÍDA")
    print("=" * 100)

if __name__ == '__main__':
    verificar_break_prevention()
