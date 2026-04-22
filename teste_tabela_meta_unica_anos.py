"""
Módulo de TESTE para tabela de meta única por ano.

Este módulo testa:
1. Carregamento dos dados do arquivo metas_institucionais_2025.xlsx
2. Geração da tabela com formato correto
3. Validação e visualização
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

from src.extractors.metas_loader import MetasInstitucionaisLoader
from src.tables.builders import adicionar_tabela_meta_unica_anos


def teste_tabela_meta_unica():
    """Teste completo da tabela de meta única por ano."""
    
    print("=" * 80)
    print("TESTE: Tabela de Meta Única por Ano")
    print("=" * 80)
    
    # ========== ETAPA 1: Carregar Dados ==========
    print("\n1️⃣  Carregando dados do arquivo metas_institucionais_2025.xlsx...")
    try:
        loader = MetasInstitucionaisLoader()
        metas = loader.obter_todas_metas()
        print(f"   ✓ {len(metas)} metas carregadas com sucesso!")
    except Exception as e:
        print(f"   ✗ Erro ao carregar: {e}")
        return False
    
    # ========== ETAPA 2: Selecionar Meta para Teste ==========
    meta_teste = "TJMG 5"
    print(f"\n2️⃣  Selecionando meta para teste: {meta_teste}")
    
    resultado = loader.obter_dados_tabela_meta_por_ano(meta_teste)
    if not resultado:
        print(f"   ✗ Meta {meta_teste} não encontrada!")
        return False
    
    anos, valores_meta, valores_resultado = resultado
    
    print(f"   ✓ Dados extraídos com sucesso!")
    print(f"   - Anos: {anos}")
    print(f"   - Valor da Meta: {valores_meta}")
    print(f"   - Resultado: {valores_resultado}")
    
    # ========== ETAPA 3: Criar Documento Word ==========
    print(f"\n3️⃣  Criando documento Word com a tabela...")
    try:
        doc = Document()
        
        # Adiciona título
        titulo = doc.add_paragraph()
        run_titulo = titulo.add_run(f"Teste de Tabela: {meta_teste}")
        run_titulo.font.size = Pt(14)
        run_titulo.font.bold = True
        
        # Adiciona descrição
        descricao = doc.add_paragraph()
        descricao.add_run(
            "Esta tabela mostra o histórico de valores e resultados de uma meta específica "
            "para cada ano disponível. A primeira linha contém o valor da meta, "
            "e a segunda linha contém o resultado apurado."
        )
        descricao.paragraph_format.space_after = Pt(12)
        
        # Adiciona tabela
        adicionar_tabela_meta_unica_anos(
            document=doc,
            nome_meta=meta_teste,
            anos=anos,
            valores_meta=valores_meta,
            valores_resultado=valores_resultado,
            titulo_custom=f"Tabela Teste - {meta_teste}: Histórico de Valores e Resultados",
            indent_cm=0,
            fonte="Fonte: Tribunal de Justiça de Minas Gerais - TJMG"
        )
        
        # Salva arquivo
        output_dir = Path(__file__).parent / "tests" / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_file = output_dir / f"teste_tabela_meta_{meta_teste.replace(' ', '_')}.docx"
        doc.save(output_file)
        
        print(f"   ✓ Documento criado com sucesso!")
        print(f"   📄 Arquivo: {output_file}")
        
    except Exception as e:
        print(f"   ✗ Erro ao criar documento: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # ========== ETAPA 4: Teste com Múltiplas Metas ==========
    print(f"\n4️⃣  Testando com múltiplas metas...")
    try:
        doc_multiplas = Document()
        
        # Adiciona título
        titulo = doc_multiplas.add_paragraph()
        run_titulo = titulo.add_run("Teste de Múltiplas Metas")
        run_titulo.font.size = Pt(14)
        run_titulo.font.bold = True
        doc_multiplas.add_paragraph()
        
        # Testa as 5 primeiras metas
        metas_teste = metas[:5]
        
        for i, meta in enumerate(metas_teste, 1):
            resultado = loader.obter_dados_tabela_meta_por_ano(meta)
            
            if resultado:
                anos, valores_meta, valores_resultado = resultado
                
                # Adiciona seção
                secao = doc_multiplas.add_paragraph()
                run_secao = secao.add_run(f"{i}. {meta}")
                run_secao.font.size = Pt(12)
                run_secao.font.bold = True
                
                # Adiciona tabela
                adicionar_tabela_meta_unica_anos(
                    document=doc_multiplas,
                    nome_meta=meta,
                    anos=anos,
                    valores_meta=valores_meta,
                    valores_resultado=valores_resultado,
                    titulo_custom=f"{meta}: Histórico de Valores",
                    fonte="Fonte: TJMG"
                )
                
                # Espaço entre tabelas
                doc_multiplas.add_paragraph()
                
                print(f"   ✓ Meta {meta}: OK")
            else:
                print(f"   ⚠ Meta {meta}: Não encontrada ou sem dados")
        
        # Salva arquivo
        output_dir = Path(__file__).parent / "tests" / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_file_multiplas = output_dir / "teste_tabela_multiplas_metas.docx"
        doc_multiplas.save(output_file_multiplas)
        
        print(f"\n   ✓ Documento com múltiplas metas criado!")
        print(f"   📄 Arquivo: {output_file_multiplas}")
        
    except Exception as e:
        print(f"   ✗ Erro ao criar documento com múltiplas metas: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # ========== ETAPA 5: Validação de Dados ==========
    print(f"\n5️⃣  Validação de dados...")
    
    # Validações
    validacoes = [
        ("Anos não vazios", len(anos) > 0),
        ("Valores da meta não vazios", len(valores_meta) > 0),
        ("Resultados não vazios", len(valores_resultado) > 0),
        ("Mesmo número de elementos", len(anos) == len(valores_meta) == len(valores_resultado)),
        ("Anos são strings", all(isinstance(a, str) for a in anos)),
        ("Primeiro valor de Meta é '—'", valores_meta[0] == '—'),
    ]
    
    validacoes_ok = 0
    for nome, resultado in validacoes:
        status = "✓" if resultado else "✗"
        print(f"   {status} {nome}: {resultado}")
        if resultado:
            validacoes_ok += 1
    
    print(f"\n   Resultado: {validacoes_ok}/{len(validacoes)} validações passaram!")
    
    # ========== Resultado Final ==========
    print("\n" + "=" * 80)
    if validacoes_ok == len(validacoes):
        print("✓ TODOS OS TESTES PASSARAM COM SUCESSO!")
        print("=" * 80)
        return True
    else:
        print("✗ ALGUNS TESTES FALHARAM!")
        print("=" * 80)
        return False


if __name__ == "__main__":
    import sys
    sucesso = teste_tabela_meta_unica()
    sys.exit(0 if sucesso else 1)
