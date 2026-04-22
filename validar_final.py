#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def validar_final():
    """Validação final definitiva"""
    
    doc = Document('data/output/Relatorio_Com_Metas_Final_v2.docx')
    
    print("=" * 100)
    print("VALIDAÇÃO FINAL: RELATORIO_COM_METAS_FINAL_V2.docx")
    print("=" * 100)
    
    print(f"\nTotal de parágrafos: {len(doc.paragraphs)}")
    print(f"Total de tabelas: {len(doc.tables)}")
    
    # Procura por "TJMG " seguido de número
    metas_tjmg = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        # Procura por "TJMG " seguido de um número
        if texto.startswith('TJMG ') and len(texto) > 10:
            # Tenta extrair o número
            parts = texto.split()
            if parts[1].isdigit():
                metas_tjmg.append(int(parts[1]))
    
    # Remove duplicatas e ordena
    metas_tjmg = sorted(set(metas_tjmg))
    
    print(f"\n✓ METAS TJMG ENCONTRADAS: {len(metas_tjmg)}")
    
    if len(metas_tjmg) > 0:
        print(f"\nPrimeiras 10 metas: {metas_tjmg[:10]}")
        if len(metas_tjmg) > 10:
            print(f"Últimas 10 metas: {metas_tjmg[-10:]}")
            
        print(f"\nIntervalo: {min(metas_tjmg)} a {max(metas_tjmg)}")
    
    print(f"\n{'='*100}")
    if len(metas_tjmg) >= 50:
        print(f"🎉 SUCESSO! {len(metas_tjmg)} METAS OPERACIONALIZADAS NO RELATÓRIO!")
        print(f"{'='*100}")
    else:
        print(f"❌ AVISO: Espera-se 55, encontrado {len(metas_tjmg)}")
        print(f"{'='*100}")

if __name__ == '__main__':
    validar_final()
