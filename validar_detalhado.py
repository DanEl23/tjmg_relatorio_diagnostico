#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def validar_metas_completo():
    """Validação mais detalhada do relatório"""
    
    doc = Document('data/output/Relatorio_Com_Metas_Final.docx')
    
    print("=" * 100)
    print("VALIDAÇÃO DETALHADA: RELATORIO_COM_METAS_FINAL.docx")
    print("=" * 100)
    
    print(f"\nTotal de parágrafos: {len(doc.paragraphs)}")
    print(f"Total de tabelas: {len(doc.tables)}")
    
    # Procura pela seção METAS INSTITUCIONAIS
    idx_secao = None
    for idx, p in enumerate(doc.paragraphs):
        if "METAS INSTITUCIONAIS" in p.text.upper():
            idx_secao = idx
            print(f"\n✓ Encontrou seção 'METAS INSTITUCIONAIS' no parágrafo {idx}")
            print(f"  Contexto: {p.text}")
            break
    
    if idx_secao is None:
        print("\n❌ Seção METAS INSTITUCIONAIS não encontrada!")
        return
    
    # Procura pelos parágrafos TJMG 5, TJMG 6, etc. APÓS a seção
    metas_encontradas = []
    for idx in range(idx_secao, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        texto = p.text.strip()
        
        # Procura por padrões como "TJMG 5", "TJMG 6", etc.
        if  texto.startswith('TJMG ') and len(texto) > 10:
            # Verifica se parece ser uma meta institucional (não orçamento)
            if '–' in texto and 'Realizar' in texto or 'Publicar' in texto or 'Reduzir' in texto or 'Implantar' in texto:
                metas_encontradas.append(texto)
                if len(metas_encontradas) <= 5:
                    print(f"\n  Meta encontrada: {texto[:70]}...")
    
    print(f"\n{'='*100}")
    print(f"✓ TOTAL DE METAS INSTITUCIONAIS ENCONTRADAS: {len(metas_encontradas)}")
    print(f"{'='*100}")
    
    if len(metas_encontradas) >= 50:
        print(f"\n🎉 SUCESSO TOTAL! {len(metas_encontradas)} metas instituciónis estão no relatório!")
    elif len(metas_encontradas) > 0:
        print(f"\n✓ Algumas metas encontradas, mas esperado 55, encontrado {len(metas_encontradas)}")
    else:
        print(f"\n❌ Nenhuma meta institucional encontrada!")
        
        # Debug: mostra os próximos 20 parágrafos após a seção
        print(f"\n--- PRÓXIMOS 20 PARÁGRAFOS APÓS 'METAS INSTITUCIONAIS': ---")
        for idx in range(idx_secao, min(idx_secao + 20, len(doc.paragraphs))):
            p = doc.paragraphs[idx]
            texto = p.text.strip()[:70]
            print(f"[{idx}] {texto}")

if __name__ == '__main__':
    validar_metas_completo()
