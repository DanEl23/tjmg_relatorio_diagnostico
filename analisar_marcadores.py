#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def analisar_marcadores():
    """Analisa como os marcadores estão formatados no Conteudo_Fonte"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("ANÁLISE DE MARCADORES EM CONTEUDO_FONTE.docx")
    print("=" * 100)
    
    # Procura por qualquer texto que pareça um marcador
    marcadores_encontrados = []
    
    for idx, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        
        # Procura por padrões que parecem marcadores
        if any(x in texto for x in ['Meta ', 'Tabela ', 'Figura ', 'Gráfico ']):
            if ':' in texto or '-' in texto or '–' in texto:
                marcadores_encontrados.append((idx, texto))
    
    print(f"\nTOTAL DE MARCADORES/RECURSOS ENCONTRADOS: {len(marcadores_encontrados)}\n")
    
    for idx, texto in marcadores_encontrados[:15]:
        print(f"[{idx}] {texto[:80]}")
    
    print("\n" + "=" * 100)
    
    # Procura especificamente pelo novo marcador
    encontrou_novo = False
    for idx, p in enumerate(doc.paragraphs):
        if "METAS_INSTITUCIONAIS" in p.text:
            print(f"\n✓ ENCONTROU NOVO MARCADOR:")
            print(f"[{idx}] {p.text}")
            encontrou_novo = True
    
    if not encontrou_novo:
        print("\n❌ NOVO MARCADOR NÃO ENCONTRADO")
        print("\nÚLTIMOS 10 PARÁGRAFOS DO DOCUMENTO:")
        for idx in range(max(0, len(doc.paragraphs) - 10), len(doc.paragraphs)):
            p = doc.paragraphs[idx]
            print(f"[{idx}] {p.text[:80]}")

if __name__ == '__main__':
    analisar_marcadores()
