#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def adicionar_marcador_metas_institucionais():
    """Adiciona o marcador MAPA_RECURSOS::METAS_INSTITUCIONAIS ao Conteudo_Fonte.docx"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("ADICIONAR MARCADOR: MAPA_RECURSOS::METAS_INSTITUCIONAIS")
    print("=" * 100)
    
    # Procura por "CONCLUSÃO" para saber onde inserir
    idx_conclusao = None
    for idx, p in enumerate(doc.paragraphs):
        if "CONCLUSÃO" in p.text and "12." in p.text:
            idx_conclusao = idx
            print(f"\n✓ Encontrado 'CONCLUSÃO' no parágrafo {idx}")
            break
    
    if idx_conclusao is None:
        print("❌ Não conseguiu encontrar a seção CONCLUSÃO")
        return
    
    # Insira antes da conclusão
    print(f"Inserindo seção de metas institucionais antes da conclusão...")
    
    # Quebra de página
    doc.paragraphs[idx_conclusao - 1].insert_paragraph_before("[QUEBRA_PAGINA]")
    
    # Título
    p_titulo = doc.paragraphs[idx_conclusao - 1].insert_paragraph_before("11. METAS INSTITUCIONAIS DO TJMG")
    p_titulo.style = 'Heading 2'
    
    # Descrição
    p_desc = doc.paragraphs[idx_conclusao - 1].insert_paragraph_before(
        "Apresentação das 55 metas institucionais do TJMG para o ano de 2025, com sua evolução "
        "comparada aos anos anteriores (2022-2024) e a meta estabelecida."
    )
    
    # Parágrafo vazio
    doc.paragraphs[idx_conclusao - 1].insert_paragraph_before("")
    
    # Marcador
    p_marcador = doc.paragraphs[idx_conclusao - 1].insert_paragraph_before("MAPA_RECURSOS::METAS_INSTITUCIONAIS")
    
    # Salva o documento
    doc.save('data/raw/Conteudo_Fonte.docx')
    
    print("✓ Marcador adicionado com sucesso!")
    print(f"✓ Documento salvo em: data/raw/Conteudo_Fonte.docx")
    print("\n" + "=" * 100)
    print("PRÓXIMO PASSO:")
    print("=" * 100)
    print("Execute o gerador de relatório com:")
    print("  python main.py")
    print("=" * 100)

if __name__ == '__main__':
    adicionar_marcador_metas_institucionais()
