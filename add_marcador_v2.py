#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def adicionar_marcador_limpo():
    """Adiciona o marcador de forma simples e correta"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("ADICIONAR MARCADOR (V2 - VERSÃO SIMPLIFICADA)")
    print("=" * 100)
    
    # Procura pela seção "12. CONCLUSÃO"
    idx_conclusao = None
    for idx, p in enumerate(doc.paragraphs):
        if "12." in p.text and "CONCLUSÃO" in p.text:
            idx_conclusao = idx
            print(f"\nEncontrou CONCLUSÃO em índice: {idx}")
            print(f"Parágrafo anterior: {doc.paragraphs[idx-1].text[:60]}")
            break
    
    if idx_conclusao is None:
        print("❌ Não encontrou CONCLUSÃO")
        return
    
    # Simplementemente adiciona os parágrafos DEPOIS do último que estiver bem formado
    # e ANTES da conclusão
    
    # Procura pelo [FINALIZAR_LISTA_MARCADORES]
    idx_final_lista = None
    for idx in range(idx_conclusao - 1, -1, -1):
        if "[FINALIZAR_LISTA_MARCADORES]" in doc.paragraphs[idx].text:
            idx_final_lista = idx
            print(f"Encontrou [FINALIZAR_LISTA_MARCADORES] em índice: {idx}")
            break
    
    if idx_final_lista is None:
        print("❌ Não encontrou [FINALIZAR_LISTA_MARCADORES]")
        return
    
    # Adiciona os parágrafos entre FINALIZAR e CONCLUSÃO
    p_anchor = doc.paragraphs[idx_final_lista]
    
    # Insere parágrafo vazio
    p_blank = p_anchor.insert_paragraph_before("")
    
    # Quebra de página (comando que o gerador processa)
    p_break = p_anchor.insert_paragraph_before("[QUEBRA_PAGINA]")
    
    # Descrição
    p_desc = p_anchor.insert_paragraph_before(
        "Apresentação das 55 metas institucionais do TJMG para o ano de 2025, com sua evolução "
        "comparada aos anos anteriores (2022-2024) e a meta estabelecida."
    )
    
    # Título
    p_titulo = p_anchor.insert_paragraph_before("11. METAS INSTITUCIONAIS DO TJMG")
    p_titulo.style = 'Heading 2'
    
    # Marcador
    p_marcador = p_anchor.insert_paragraph_before("MAPA_RECURSOS::METAS_INSTITUCIONAIS")
    
    # Salva
    doc.save('data/raw/Conteudo_Fonte.docx')
    
    print("\n✓ Marcador adicionado com sucesso!")
    print("✓ Posição: Após [FINALIZAR_LISTA_MARCADORES] e antes de CONCLUSÃO")
    print("✓ Documento salvo!")

if __name__ == '__main__':
    adicionar_marcador_limpo()
