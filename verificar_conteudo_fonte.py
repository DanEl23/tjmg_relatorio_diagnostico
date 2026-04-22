#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

def verificar_conteudo_fonte():
    """Verifica se Conteudo_Fonte.docx contém o marcador METAS_INSTITUCIONAIS"""
    
    doc = Document('data/raw/Conteudo_Fonte.docx')
    
    print("=" * 100)
    print("VERIFICAÇÃO DE CONTEUDO_FONTE.docx")
    print("=" * 100)
    print(f"\nTotal de parágrafos: {len(doc.paragraphs)}\n")
    
    encontrou_metas = False
    
    for idx, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if "METAS_INSTITUCIONAIS" in texto or "Metas Institucionais" in texto.lower():
            encontrou_metas = True
            print(f"[Parágrafo {idx}]: {texto}")
    
    if not encontrou_metas:
        print("❌ NÃO ENCONTRADO: Marcador 'METAS_INSTITUCIONAIS' ou 'Metas Institucionais'")
    else:
        print("\n✓ ENCONTRADO: Marcador para metas institucionais")
    
    # Mostra os últimos parágrafos para contextualizar onde adicionar
    print("\n" + "=" * 100)
    print("ÚLTIMOS 15 PARÁGRAFOS DO DOCUMENTO:")
    print("=" * 100)
    start_idx = max(0, len(doc.paragraphs) - 15)
    for idx in range(start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        texto = p.text.strip()
        if texto:
            print(f"[{idx}] {texto[:80]}...")
        else:
            print(f"[{idx}] (parágrafo vazio)")

if __name__ == '__main__':
    verificar_conteudo_fonte()
